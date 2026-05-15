import json
import os
import re
import sys
import tempfile
import threading
import time
import uuid
import webbrowser
import zipfile
from copy import deepcopy
from difflib import SequenceMatcher
from io import BytesIO
from pathlib import Path

import fitz  # PyMuPDF
import pytesseract
from google import genai
from PIL import Image, ImageDraw, ImageFilter, ImageFont
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from flask import Flask, jsonify, redirect, render_template, request, send_file, send_from_directory

# ---------------------------------------------------------------------------
# PyInstaller-aware paths
# ---------------------------------------------------------------------------
_FROZEN   = getattr(sys, "frozen", False)
_BUNDLE   = Path(sys._MEIPASS) if _FROZEN else Path(__file__).parent
_EXE_DIR  = Path(sys.executable).parent if _FROZEN else Path(__file__).parent

# Fix SSL certificate lookup when running as a frozen exe
if _FROZEN:
    import certifi
    os.environ.setdefault("SSL_CERT_FILE",      certifi.where())
    os.environ.setdefault("REQUESTS_CA_BUNDLE", certifi.where())

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
if _FROZEN:
    UPLOAD_FOLDER = _EXE_DIR / "data" / "uploads"
    OUTPUT_FOLDER = _EXE_DIR / "data" / "outputs"
else:
    UPLOAD_FOLDER = Path(tempfile.gettempdir()) / "butterlayer" / "uploads"
    OUTPUT_FOLDER = Path(tempfile.gettempdir()) / "butterlayer" / "outputs"
UPLOAD_FOLDER.mkdir(parents=True, exist_ok=True)
OUTPUT_FOLDER.mkdir(parents=True, exist_ok=True)

# Tesseract: bundled copy next to exe takes priority over system install
_BUNDLED_TESS = _EXE_DIR / "tesseract" / "tesseract.exe"
if _BUNDLED_TESS.exists():
    TESSERACT_CMD = str(_BUNDLED_TESS)
    os.environ.setdefault("TESSDATA_PREFIX", str(_EXE_DIR / "tesseract" / "tessdata"))
else:
    TESSERACT_CMD = os.environ.get(
        "TESSERACT_CMD",
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    )
if Path(TESSERACT_CMD).exists():
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")
TARGET_LANGUAGE = os.environ.get("TARGET_LANGUAGE", "Traditional Chinese (繁體中文)")
GEMINI_MODEL = "gemini-2.5-flash"

def _create_desktop_shortcut():
    if not _FROZEN:
        return
    try:
        import subprocess
        exe = str(Path(sys.executable))
        work_dir = str(Path(sys.executable).parent)
        ps = (
            '$desktop=[Environment]::GetFolderPath("Desktop");'
            '$lnk="$desktop\\ButterLayer.lnk";'
            'if(Test-Path $lnk){exit 0};'
            '$s=(New-Object -COM WScript.Shell).CreateShortcut($lnk);'
            f'$s.TargetPath="{exe}";'
            f'$s.WorkingDirectory="{work_dir}";'
            '$s.Save()'
        )
        subprocess.run(
            ["powershell", "-NoProfile", "-NonInteractive", "-Command", ps],
            capture_output=True,
            creationflags=subprocess.CREATE_NO_WINDOW,
        )
    except Exception:
        pass
GEMINI_FALLBACK_MODELS = ["gemini-2.5-flash-lite", "gemini-2.5-flash", "gemini-2.5-pro"]
PAGE_DELAY_SECONDS = 0
_dead_models: set[str] = set()        # 404 ??permanently unavailable
_cooling_models: dict[str, float] = {}  # 503 exhausted ??cooldown until timestamp


def _is_cooling(model: str) -> bool:
    exp = _cooling_models.get(model)
    if exp is None:
        return False
    if time.time() < exp:
        return True
    del _cooling_models[model]
    return False


def _model_cost_rank(model_id: str) -> float:
    """Lower = cheaper. Ranks by version number then tier (lite < flash < pro)."""
    import re
    m = re.search(r'(\d+[\.,]\d+)', model_id)
    version = float(m.group(1).replace(',', '.')) if m else 99.0
    tier = 1 if 'lite' in model_id else (3 if 'pro' in model_id else 2)
    return version * 10 + tier

# ---------------------------------------------------------------------------
# In-memory stores
# ---------------------------------------------------------------------------
jobs: dict[str, dict] = {}
job_results: dict[str, dict] = {}
STORE_LOCK = threading.Lock()
TRANSLATION_BLOCK_STEP = 1000.0
ALLOWED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}
NATIVE_IMAGE_RENDER_DPI = 360
NATIVE_IMAGE_MAX_PIXELS = 36_000_000
INSTANCE_LOCK_HANDLE = None

app = Flask(__name__, template_folder=str(_BUNDLE / "templates"))
app.config["MAX_CONTENT_LENGTH"] = 500 * 1024 * 1024  # 500 MB


@app.errorhandler(413)
def _too_large(e):
    return jsonify({"error": "檔案太大，請上傳 500 MB 以下的 PDF。"}), 413

@app.errorhandler(Exception)
def _json_error(e):
    import traceback
    return jsonify({"error": str(e), "detail": traceback.format_exc()}), 500


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_job(job_id: str, **kwargs):
    with STORE_LOCK:
        if job_id not in jobs:
            jobs[job_id] = {}
        jobs[job_id].update(kwargs)


def _get_job(job_id: str) -> dict | None:
    with STORE_LOCK:
        job = jobs.get(job_id)
        return dict(job) if job else None


def _get_result(job_id: str) -> dict | None:
    with STORE_LOCK:
        data = job_results.get(job_id)
    if data is None:
        # Try loading from disk (server may have restarted since job completed)
        path = OUTPUT_FOLDER / f"{job_id}_result.json"
        if path.exists():
            try:
                with open(str(path), encoding="utf-8") as f:
                    data = json.load(f)
                with STORE_LOCK:
                    job_results[job_id] = data
            except Exception:
                return None
    if data is None:
        return None
    raw_data = deepcopy(data)
    _repair_scanned_overlay_data(raw_data)
    return _serialize_result_data(job_id, raw_data)


def _normalize_text(text: str) -> str:
    t = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    # Strip characters illegal in XML (NULL bytes, most C0/C1 control chars)
    # Keep: tab (0x09), newline (0x0A), carriage return (0x0D), and 0x20+
    return "".join(
        ch for ch in t
        if ch in "\t\n\r" or "\x20" <= ch <= "\ud7ff"
        or "\ue000" <= ch <= "\ufffd" or "\U00010000" <= ch <= "\U0010ffff"
    )


def _save_partial_result(job_id: str, data: dict) -> None:
    try:
        path = OUTPUT_FOLDER / f"{job_id}_result.json"
        OUTPUT_FOLDER.mkdir(parents=True, exist_ok=True)
        with open(str(path), "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False)
    except Exception as e:
        print(f"[partial save] failed: {e}")


def _load_saved_results() -> None:
    try:
        for path in OUTPUT_FOLDER.glob("*_result.json"):
            job_id = path.stem.replace("_result", "")
            if job_id in job_results:
                continue
            with open(str(path), "r", encoding="utf-8") as f:
                data = json.load(f)
            if _repair_scanned_overlay_data(data):
                _save_partial_result(job_id, data)
            with STORE_LOCK:
                job_results[job_id] = data
            print(f"[startup] loaded saved result: {data.get('source_name', job_id)}")
    except Exception as e:
        print(f"[startup] failed to load saved results: {e}")


def _serialize_result_data(job_id: str, data: dict) -> dict:
    page_lines = [[dict(line) for line in page] for page in data.get("page_lines", [])]
    page_words = [[dict(w) for w in page] for page in data.get("page_words", [])]
    page_source_blocks = [[dict(b) for b in page] for page in data.get("page_source_blocks", [])]
    page_ocr_paras = [[dict(p) for p in page] for page in data.get("page_ocr_paras", [])]
    return {
        "job_id": job_id,
        "original": list(data.get("original", [])),
        "translated": list(data.get("translated", [])),
        "page_lines": page_lines,
        "page_words": page_words,
        "page_source_blocks": page_source_blocks,
        "page_ocr_paras": page_ocr_paras,
        "translated_blocks": [
            [_serialize_translation_block(job_id, block) for block in page]
            for page in data.get("translated_blocks", [])
        ],
        "translated_blueprints": [
            [_serialize_translation_block(job_id, block) for block in page]
            for page in data.get("translated_blueprints", [])
        ],
        "target_lang": data.get("target_lang", TARGET_LANGUAGE),
        "source_name": data.get("source_name", ""),
        "updated_at": data.get("updated_at"),
    }


def _normalize_line_box(x0: float, y0: float, x1: float, y1: float, width: float, height: float) -> dict | None:
    width = max(float(width), 1.0)
    height = max(float(height), 1.0)
    if x1 <= x0 or y1 <= y0:
        return None

    return {
        "left": round(max(0.0, min(1.0, x0 / width)), 6),
        "right": round(max(0.0, min(1.0, x1 / width)), 6),
        "top": round(max(0.0, min(1.0, y0 / height)), 6),
        "bottom": round(max(0.0, min(1.0, y1 / height)), 6),
    }


def _text_to_translation_blocks(text: str) -> list[dict]:
    lines = _normalize_text(text).split("\n")
    return [
        {"type": "text", "source_idx": round(i * TRANSLATION_BLOCK_STEP, 6), "text": line}
        for i, line in enumerate(lines)
    ]


def _pages_to_translation_blocks(pages: list[str]) -> list[list[dict]]:
    return [_text_to_translation_blocks(page) for page in pages]


def _split_translation_text_for_paras(text: str, para_count: int) -> list[str]:
    text = _normalize_text(text).strip()
    if para_count <= 0:
        return []
    if para_count == 1:
        return [text]

    paras = [p.strip() for p in re.split(r"\n\s*\n+", text) if p.strip()]
    if len(paras) == para_count:
        return paras

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if len(lines) == para_count:
        return lines

    chunks = paras or lines or ([text] if text else [])
    if len(chunks) >= para_count:
        return chunks[: para_count - 1] + ["\n".join(chunks[para_count - 1:])]
    return chunks + [""] * (para_count - len(chunks))


def _has_para_idx_blocks(blocks: object) -> bool:
    if not isinstance(blocks, list):
        return False
    return any(
        isinstance(block, dict)
        and block.get("type", "text") != "image"
        and block.get("para_idx") is not None
        for block in blocks
    )


def _repair_scanned_overlay_data(data: dict) -> bool:
    """Backfill OCR paragraph mapping for scanned/old partial results.

    Older partial saves flattened scanned translations into plain text blocks and
    dropped page_ocr_paras.  That makes the DOM overlay unable to place text
    after a restart.  Rebuild the mapping from saved OCR line boxes when possible.
    """
    source_blocks = data.get("page_source_blocks")
    if isinstance(source_blocks, list) and any(isinstance(page, list) and page for page in source_blocks):
        return False

    translated = data.get("translated", [])
    page_lines = data.get("page_lines", [])
    if not isinstance(translated, list) or not isinstance(page_lines, list) or not page_lines:
        return False

    page_ocr_paras = data.get("page_ocr_paras")
    if not isinstance(page_ocr_paras, list):
        page_ocr_paras = []

    translated_blocks = data.get("translated_blocks")
    if not isinstance(translated_blocks, list):
        translated_blocks = _pages_to_translation_blocks(translated)

    total_pages = max(len(translated), len(page_lines), len(page_ocr_paras), len(translated_blocks))
    changed = False

    while len(page_ocr_paras) < total_pages:
        page_ocr_paras.append([])
        changed = True
    while len(translated_blocks) < total_pages:
        fallback_text = translated[len(translated_blocks)] if len(translated_blocks) < len(translated) else ""
        translated_blocks.append(_text_to_translation_blocks(fallback_text))
        changed = True

    for page_index in range(total_pages):
        lines = page_lines[page_index] if page_index < len(page_lines) else []
        if not isinstance(lines, list) or not lines:
            continue

        paras = page_ocr_paras[page_index] if page_index < len(page_ocr_paras) else []
        if not isinstance(paras, list) or not paras:
            paras = _group_ocr_paragraphs(lines)
            if paras:
                page_ocr_paras[page_index] = paras
                changed = True
        if not paras:
            continue

        page_blocks = translated_blocks[page_index]
        if _has_para_idx_blocks(page_blocks):
            continue

        page_text = translated[page_index] if page_index < len(translated) else _translation_blocks_to_text(page_blocks)
        chunks = _split_translation_text_for_paras(page_text, len(paras))
        image_blocks = [
            dict(block)
            for block in page_blocks
            if isinstance(block, dict) and block.get("type") == "image"
        ]
        repaired = [
            {
                "type": "text",
                "para_idx": pi,
                "source_idx": round(pi * TRANSLATION_BLOCK_STEP, 6),
                "text": chunks[pi] if pi < len(chunks) else "",
            }
            for pi in range(len(paras))
        ]
        repaired.extend(image_blocks)
        translated_blocks[page_index] = sorted(
            repaired,
            key=lambda block: float(block.get("source_idx", 0) or 0),
        )
        changed = True

    if changed:
        data["page_ocr_paras"] = page_ocr_paras
        data["translated_blocks"] = translated_blocks
        blueprints = data.get("translated_blueprints")
        if not isinstance(blueprints, list) or not any(_has_para_idx_blocks(page) for page in blueprints):
            data["translated_blueprints"] = _clone_translation_blocks(translated_blocks)
    return changed


def _clone_translation_blocks(pages: list[list[dict]]) -> list[list[dict]]:
    return [[dict(block) for block in page] for page in pages]


def _get_job_asset_dir(job_id: str) -> Path:
    asset_dir = OUTPUT_FOLDER / f"{job_id}_assets"
    asset_dir.mkdir(parents=True, exist_ok=True)
    return asset_dir


def _job_image_path(job_id: str, image_id: str) -> Path:
    safe_name = Path(image_id).name
    return _get_job_asset_dir(job_id) / safe_name


def _job_image_url(job_id: str, image_id: str) -> str:
    return f"/result/{job_id}/image/{Path(image_id).name}"


def _serialize_translation_block(job_id: str, block: dict) -> dict:
    data = dict(block)
    block_type = data.get("type", "text")
    if block_type == "image" and data.get("image_id"):
        data["image_url"] = _job_image_url(job_id, str(data["image_id"]))
    return data


def _normalize_translation_blocks(blocks: object, fallback_text: str = "", job_id: str | None = None) -> list[dict]:
    if not isinstance(blocks, list):
        return _text_to_translation_blocks(fallback_text)

    normalized = []
    for i, block in enumerate(blocks):
        if not isinstance(block, dict):
            continue

        try:
            source_idx = float(block.get("source_idx", i * TRANSLATION_BLOCK_STEP))
        except (TypeError, ValueError):
            source_idx = i * TRANSLATION_BLOCK_STEP

        block_type = str(block.get("type", "text")).strip().lower() or "text"
        if block_type == "image":
            image_id = str(block.get("image_id", "")).strip()
            if not image_id:
                continue
            if job_id is not None and not _job_image_path(job_id, image_id).exists():
                continue
            normalized.append({
                "type": "image",
                "source_idx": round(source_idx, 6),
                "image_id": Path(image_id).name,
                "name": _normalize_text(str(block.get("name", "")).strip()),
                "width": max(int(block.get("width", 0) or 0), 0),
                "height": max(int(block.get("height", 0) or 0), 0),
            })
            continue

        try:
            blk_font_size = max(8, min(72, int(block.get("font_size") or 10)))
        except (TypeError, ValueError):
            blk_font_size = 10
        text_block: dict = {
            "type": "text",
            "source_idx": round(source_idx, 6),
            "text": _normalize_text(str(block.get("text", ""))),
            "font_size": blk_font_size,
            "bold": bool(block.get("bold", False)),
        }
        if block.get("para_idx") is not None:
            text_block["para_idx"] = int(block["para_idx"])
        normalized.append(text_block)

    if not normalized:
        return _text_to_translation_blocks(fallback_text)

    return normalized


def _translation_blocks_to_text(blocks: list[dict]) -> str:
    return "\n".join(
        _normalize_text(str(block.get("text", "")))
        for block in blocks
        if block.get("type", "text") != "image"
    )


def _iter_translation_blocks(
    translated_pages: list[str],
    translated_blocks: list[list[dict]] | None,
):
    pages = translated_blocks if translated_blocks is not None else _pages_to_translation_blocks(translated_pages)
    for page in pages:
        for block in page:
            yield block


def _looks_like_title(text: str, next_text: str) -> bool:
    clean_text = _normalize_text(text).strip()
    clean_next = _normalize_text(next_text).strip()
    if not clean_text or not clean_next:
        return False
    return len(clean_text) <= 80 and len(clean_next) >= 80


def _apply_reading_doc_style(doc: Document, font_size: int = 10):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Segoe UI"
    normal_style.font.size = Pt(font_size)


def _style_run(run, size=None, bold: bool | None = None):
    run.font.name = "Segoe UI"
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold


def _section_content_size_emu(section) -> tuple[int, int]:
    width = int(section.page_width - section.left_margin - section.right_margin)
    height = int(section.page_height - section.top_margin - section.bottom_margin)
    return max(1, width), max(1, height)


def _image_aspect_ratio(image_source) -> float | None:
    try:
        if hasattr(image_source, "seek"):
            image_source.seek(0)
        with Image.open(image_source) as img:
            w, h = img.size
        if hasattr(image_source, "seek"):
            image_source.seek(0)
        if w > 0 and h > 0:
            return float(w) / float(h)
    except Exception:
        return None
    return None


def _add_picture_within_bounds(
    run,
    image_source,
    max_width_emu: int,
    max_height_emu: int,
    safety_pt: float = 10.0,
):
    aspect = _image_aspect_ratio(image_source)
    if aspect is None or aspect <= 0:
        run.add_picture(image_source, width=max(1, int(max_width_emu)))
        return

    safety_emu = int(Pt(max(0.0, safety_pt)))
    usable_w = max(1, int(max_width_emu) - safety_emu)
    usable_h = max(1, int(max_height_emu) - safety_emu)
    target_w = usable_w
    target_h = int(round(target_w / aspect))
    if target_h > usable_h:
        target_h = usable_h
        target_w = int(round(target_h * aspect))
    run.add_picture(
        image_source,
        width=max(1, target_w),
        height=max(1, target_h),
    )


def _get_client(api_key: str) -> genai.Client:
    return genai.Client(api_key=api_key)


def _acquire_instance_lock() -> bool:
    """Prevent multiple local ButterLayer servers from serving different frontend builds."""
    global INSTANCE_LOCK_HANDLE
    lock_dir = Path(tempfile.gettempdir()) / "butterlayer"
    lock_dir.mkdir(parents=True, exist_ok=True)
    lock_path = lock_dir / "butterlayer_5000.lock"
    try:
        import msvcrt
        handle = open(str(lock_path), "a+b")
        handle.seek(0)
        msvcrt.locking(handle.fileno(), msvcrt.LK_NBLCK, 1)
        INSTANCE_LOCK_HANDLE = handle
        return True
    except Exception:
        try:
            handle.close()  # type: ignore[name-defined]
        except Exception:
            pass
        return False


def _extract_native_page_lines(page: fitz.Page) -> list[dict]:
    page_dict = page.get_text("dict")
    lines = []
    for block in page_dict.get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            bbox = line.get("bbox")
            spans = line.get("spans", [])
            text = "".join(span.get("text", "") for span in spans).strip()
            if not bbox or not text:
                continue
            box = _normalize_line_box(*bbox, page.rect.width, page.rect.height)
            if box:
                lines.append(box)

    return sorted(lines, key=lambda item: (item["top"], item["left"]))


def _extract_native_page_words(page: fitz.Page) -> list[dict]:
    W, H = page.rect.width, page.rect.height
    words = []
    for w in page.get_text("words"):
        x0, y0, x1, y1, text = w[0], w[1], w[2], w[3], str(w[4]).strip()
        if not text:
            continue
        box = _normalize_line_box(x0, y0, x1, y1, W, H)
        if box:
            box["text"] = text
            words.append(box)
    return sorted(words, key=lambda b: (b["top"], b["left"]))


def _find_table_rects(page: fitz.Page) -> list:
    """Detect table regions on a native-text PDF page using PyMuPDF find_tables."""
    try:
        finder = page.find_tables()
        return [fitz.Rect(t.bbox) for t in (finder.tables or [])]
    except Exception:
        return []


def _extract_figure_images(
    page: fitz.Page, job_id: str, page_num: int
) -> tuple[list[dict], list]:
    """Detect embedded raster image blocks (PDF block type==1) and render them.
    Returns (figure_block_infos, exclude_rects) ??exclude_rects covers each figure
    plus a margin so nearby axis labels / tick marks are also skipped."""
    if not job_id:
        return [], []
    raw_blocks = page.get_text("blocks")
    img_blocks = [b for b in raw_blocks if b[6] == 1]
    if not img_blocks:
        return [], []

    page_rect = page.rect
    page_h = max(page_rect.height, 1)
    infos: list[dict] = []
    exclude: list = []

    for i, b in enumerate(img_blocks):
        rect = fitz.Rect(b[0], b[1], b[2], b[3])
        clip = fitz.Rect(
            max(rect.x0 - 4, page_rect.x0), max(rect.y0 - 4, page_rect.y0),
            min(rect.x1 + 4, page_rect.x1), min(rect.y1 + 4, page_rect.y1),
        )
        pix = page.get_pixmap(matrix=mat, clip=clip)
        image_id = f"fig_p{page_num}_i{i}.png"
        _job_image_path(job_id, image_id).parent.mkdir(parents=True, exist_ok=True)
        pix.save(str(_job_image_path(job_id, image_id)))
        infos.append({"image_id": image_id, "y_frac": rect.y0 / page_h})
        # Expand exclusion zone to swallow axis labels / caption text around the figure
        margin = min(max(rect.height * 0.25, 20), 50)
        exclude.append(fitz.Rect(
            rect.x0 - margin, rect.y0 - margin,
            rect.x1 + margin, rect.y1 + margin,
        ))

    return infos, exclude


def _rect_overlap_ratio(a: fitz.Rect, b: fitz.Rect) -> float:
    ix0, iy0 = max(a.x0, b.x0), max(a.y0, b.y0)
    ix1, iy1 = min(a.x1, b.x1), min(a.y1, b.y1)
    if ix0 >= ix1 or iy0 >= iy1:
        return 0.0
    return ((ix1 - ix0) * (iy1 - iy0)) / max(a.get_area(), 1.0)


def _rect_overlap_min_ratio(a: fitz.Rect, b: fitz.Rect) -> float:
    ix0, iy0 = max(a.x0, b.x0), max(a.y0, b.y0)
    ix1, iy1 = min(a.x1, b.x1), min(a.y1, b.y1)
    if ix0 >= ix1 or iy0 >= iy1:
        return 0.0
    return ((ix1 - ix0) * (iy1 - iy0)) / max(min(a.get_area(), b.get_area()), 1.0)


def _rect_has_meaningful_overlap(a: fitz.Rect, b: fitz.Rect) -> bool:
    ix0, iy0 = max(a.x0, b.x0), max(a.y0, b.y0)
    ix1, iy1 = min(a.x1, b.x1), min(a.y1, b.y1)
    if ix0 >= ix1 or iy0 >= iy1:
        return False
    overlap_w = ix1 - ix0
    overlap_h = iy1 - iy0
    return (
        overlap_w >= min(10.0, max(a.width, 1.0) * 0.15)
        and overlap_h >= min(5.0, max(a.height, 1.0) * 0.15)
    )


def _expanded_clip(rect: fitz.Rect, page_rect: fitz.Rect, margin: float = 4) -> fitz.Rect:
    return fitz.Rect(
        max(rect.x0 - margin, page_rect.x0),
        max(rect.y0 - margin, page_rect.y0),
        min(rect.x1 + margin, page_rect.x1),
        min(rect.y1 + margin, page_rect.y1),
    )


def _native_image_matrix_for_clip(clip: fitz.Rect) -> fitz.Matrix:
    """Higher-resolution render for native-PDF figure extraction.

    Keeps a pixel ceiling to avoid excessive memory usage on large clips.
    """
    width = max(float(clip.width), 1.0)
    height = max(float(clip.height), 1.0)
    target_scale = max(2.0, NATIVE_IMAGE_RENDER_DPI / 72.0)
    max_scale = (NATIVE_IMAGE_MAX_PIXELS / (width * height)) ** 0.5
    scale = max(2.0, min(target_scale, max_scale))
    return fitz.Matrix(scale, scale)


def _union_rect(rects: list[fitz.Rect]) -> fitz.Rect | None:
    if not rects:
        return None
    merged = fitz.Rect(rects[0])
    for rect in rects[1:]:
        merged.include_rect(rect)
    return merged


def _collect_graphic_rects(page: fitz.Page, table_rects: list | None = None) -> list[fitz.Rect]:
    """Collect visible non-text areas, including raster images and vector drawings."""
    rects: list[fitz.Rect] = [fitz.Rect(r) for r in (table_rects or [])]
    seen_xrefs: set[int] = set()

    for image in page.get_images(full=True):
        xref = int(image[0])
        if xref in seen_xrefs:
            continue
        seen_xrefs.add(xref)
        try:
            rects.extend(fitz.Rect(r) for r in page.get_image_rects(xref))
        except Exception:
            continue
    try:
        for drawing in page.get_drawings():
            rect = drawing.get("rect")
            if rect:
                frect = fitz.Rect(rect)
                if frect.get_area() > 10:
                    rects.append(frect)
    except Exception:
        pass

    page_rect = page.rect
    page_area = page_rect.get_area()
    return [
        r for r in rects
        if r.get_area() > 10
        and r.get_area() < page_area * 0.80  # exclude full-page backgrounds (scanned pages)
        and r.width > 2
        and r.height > 2
        and r.intersects(page_rect)
    ]


def _page_has_fullpage_raster_background(page: fitz.Page, min_area_ratio: float = 0.90) -> bool:
    page_area = max(page.rect.get_area(), 1.0)
    try:
        for image in page.get_images(full=True):
            xref = int(image[0])
            try:
                rects = [fitz.Rect(r) for r in page.get_image_rects(xref)]
            except Exception:
                continue
            if any(rect.get_area() >= page_area * min_area_ratio for rect in rects):
                return True
    except Exception:
        return False
    return False


_CAPTION_RE = re.compile(
    r"^\s*(?:(?:Fig(?:ure)?\.?)|(?:Table))\s*\d+[A-Za-z]?\b",
    re.IGNORECASE,
)


def _is_visual_caption(text: str) -> bool:
    normalized = " ".join(_normalize_text(text).strip().split())
    match = _CAPTION_RE.search(normalized)
    if not match:
        return False
    after = normalized[match.end():].lstrip(" .:-").lower()
    if after.startswith((
        "shows ",
        "show ",
        "summarizes ",
        "summarises ",
        "lists ",
        "presents ",
        "indicates ",
        "illustrates ",
        "compares ",
    )):
        return False
    return True


def _page_has_side_by_side_columns(page: fitz.Page, text_blocks: list | None = None) -> bool:
    blocks = text_blocks
    if blocks is None:
        blocks = [
            b for b in page.get_text("blocks")
            if len(b) >= 7 and b[6] == 0 and str(b[4]).strip()
        ]
    page_w = max(page.rect.width, 1)
    narrative = []
    for b in blocks:
        text = str(b[4]).strip()
        clean = " ".join(_normalize_text(text).split())
        if len(clean) < 45:
            continue
        if _is_formula_text(text) or _is_formula_adjunct_text(text):
            continue
        narrative.append(b)
    narrow = [b for b in narrative if (b[2] - b[0]) <= page_w * 0.55]
    left = [b for b in narrow if ((b[0] + b[2]) / 2) < page_w * 0.48]
    right = [b for b in narrow if ((b[0] + b[2]) / 2) > page_w * 0.52]
    if len(left) < 2 or len(right) < 2:
        return False
    left_y = (min(b[1] for b in left), max(b[3] for b in left))
    right_y = (min(b[1] for b in right), max(b[3] for b in right))
    overlap = min(left_y[1], right_y[1]) - max(left_y[0], right_y[0])
    return overlap > page.rect.height * 0.20


def _visual_order_key(rect: fitz.Rect, page: fitz.Page, text_blocks: list | None = None) -> float:
    """Return a 0..1 key that follows reading order, not only vertical position."""
    page_w = max(page.rect.width, 1)
    page_h = max(page.rect.height, 1)
    y_key = max(0.0, min(1.0, rect.y0 / page_h))
    if not _page_has_side_by_side_columns(page, text_blocks):
        return y_key
    if rect.width > page_w * 0.55:
        return y_key
    col = 0 if ((rect.x0 + rect.x1) / 2) < page_w * 0.5 else 1
    return (col + y_key) / 2


def _same_reading_column(a: fitz.Rect, b: fitz.Rect, page_w: float) -> bool:
    """True when two narrow regions belong to the same visual reading column."""
    if a.width > page_w * 0.55 or b.width > page_w * 0.55:
        return True
    a_left = ((a.x0 + a.x1) / 2) < page_w * 0.5
    b_left = ((b.x0 + b.x1) / 2) < page_w * 0.5
    return a_left == b_left


def _estimate_source_idx_from_text_blocks(
    page: fitz.Page,
    rect: fitz.Rect,
    text_blocks: list | None = None,
    exclude_rects: list | None = None,
) -> float:
    blocks = text_blocks
    if blocks is None:
        blocks = [
            b for b in page.get_text("blocks")
            if len(b) >= 7 and b[6] == 0 and str(b[4]).strip()
        ]
    rect_key = _visual_order_key(rect, page, blocks)
    before = 0
    for block in blocks:
        block_rect = fitz.Rect(block[0], block[1], block[2], block[3])
        if _rect_overlap_ratio(block_rect, rect) > 0.35:
            continue
        if exclude_rects and any(_rect_overlap_ratio(block_rect, er) > 0.35 for er in exclude_rects):
            continue
        if _visual_order_key(block_rect, page, blocks) < rect_key:
            block_text = _normalize_text(str(block[4])).strip()
            line_count = len(block_text.splitlines())
            before += max(1, line_count) + 1
    return round(before * TRANSLATION_BLOCK_STEP - TRANSLATION_BLOCK_STEP / 2, 6)


def _filtered_text_blocks(page: fitz.Page, exclude_rects: list | None = None) -> list:
    """Return text blocks with only the spans outside excluded image/formula clips."""
    if not exclude_rects:
        return page.get_text("blocks")

    ex_rects = [fitz.Rect(r) for r in exclude_rects]

    def _is_excluded_rect(candidate: fitz.Rect) -> bool:
        if candidate.is_empty:
            return True
        area = max(candidate.get_area(), 1.0)
        cx = (candidate.x0 + candidate.x1) / 2
        cy = (candidate.y0 + candidate.y1) / 2
        for er in ex_rects:
            if er.x0 <= cx <= er.x1 and er.y0 <= cy <= er.y1:
                return True
            inter = candidate & er
            if inter.is_empty:
                continue
            inter_area = inter.get_area()
            if inter_area / area > 0.2:
                return True
            if inter_area / max(min(candidate.get_area(), er.get_area()), 1.0) > 0.45:
                return True
        return False

    raw = page.get_text("rawdict")
    blocks: list = []
    for block_no, block in enumerate(raw.get("blocks", [])):
        if block.get("type") != 0:
            continue
        kept_lines: list[str] = []
        kept_rects: list[fitz.Rect] = []
        for line in block.get("lines", []):
            line_parts: list[str] = []
            for span in line.get("spans", []):
                chars = span.get("chars", [])
                text = "".join(ch.get("c", "") for ch in chars) if chars else span.get("text", "")
                if not text:
                    continue
                span_rect = fitz.Rect(span.get("bbox", line.get("bbox", block.get("bbox"))))
                if _is_excluded_rect(span_rect):
                    continue
                line_parts.append(text)
                kept_rects.append(span_rect)
            if line_parts:
                line_text = "".join(line_parts)
                line_text = re.sub(r"^\s*\)\s*", "", line_text)
                kept_lines.append(line_text)
        text = "\n".join(kept_lines).strip()
        if not text:
            continue
        rect = _union_rect(kept_rects) or fitz.Rect(block.get("bbox"))
        if _is_excluded_rect(rect):
            continue
        blocks.append((rect.x0, rect.y0, rect.x1, rect.y1, text, block_no, 0))
    return blocks


def _extract_captioned_visuals(
    page: fitz.Page, job_id: str, page_num: int, table_rects: list | None = None
) -> tuple[list[dict], list]:
    """Render captioned figures/tables as images and exclude them from translation."""
    if not job_id:
        return [], []

    text_blocks = [
        b for b in page.get_text("blocks")
        if len(b) >= 7 and b[6] == 0 and str(b[4]).strip()
    ]
    captions = [
        b for b in text_blocks
        if _is_visual_caption(str(b[4]))
    ]
    if not captions:
        return [], []

    page_rect = page.rect
    page_w = max(page_rect.width, 1)
    page_h = max(page_rect.height, 1)
    graphics = _collect_graphic_rects(page, table_rects=table_rects)
    infos: list[dict] = []
    exclude: list = []

    def _connected_above_graphics(items: list[tuple[float, fitz.Rect]]) -> list[fitz.Rect]:
        if not items:
            return []
        min_gap = min(gap for gap, _ in items)
        selected = [
            graphic for gap, graphic in items
            if gap <= min_gap + 8
        ]
        max_stack_gap = max(34.0, page_h * 0.055)
        while True:
            current = _union_rect(selected)
            if current is None:
                return selected
            added = False
            for _, graphic in sorted(items, key=lambda item: item[1].y1, reverse=True):
                if any(_rect_overlap_ratio(graphic, used) > 0.85 for used in selected):
                    continue
                if graphic.y0 >= current.y0:
                    continue
                vertical_gap = max(0.0, current.y0 - graphic.y1)
                horizontal_overlap = min(current.x1, graphic.x1) - max(current.x0, graphic.x0)
                if vertical_gap <= max_stack_gap and horizontal_overlap > -page_w * 0.12:
                    selected.append(graphic)
                    added = True
                    break
            if not added:
                return selected

    for i, caption in enumerate(captions):
        cap_rect = fitz.Rect(caption[0], caption[1], caption[2], caption[3])
        caption_text = " ".join(_normalize_text(str(caption[4])).strip().split())
        is_table_caption = caption_text.lower().startswith("table")
        same_row_graphics: list[fitz.Rect] = []
        above_graphics: list[tuple[float, fitz.Rect]] = []
        below_graphics: list[tuple[float, fitz.Rect]] = []
        for graphic in graphics:
            if any(_rect_overlap_ratio(graphic, used_rect) > 0.5 for used_rect in exclude):
                continue
            vertical_overlap = min(cap_rect.y1, graphic.y1) - max(cap_rect.y0, graphic.y0)
            horizontal_overlap = min(cap_rect.x1, graphic.x1) - max(cap_rect.x0, graphic.x0)
            caption_overlap_ratio = (
                (cap_rect & graphic).get_area() / max(cap_rect.get_area(), 1.0)
                if cap_rect.intersects(graphic) else 0.0
            )
            cap_center_x = (cap_rect.x0 + cap_rect.x1) / 2
            graphic_center_x = (graphic.x0 + graphic.x1) / 2
            same_half = (
                graphic_center_x < page_w * 0.52
                if cap_center_x < page_w * 0.5
                else graphic_center_x > page_w * 0.48
            )
            adjacent_same_band = (
                graphic.x0 >= cap_rect.x1 - 8
                and graphic.x0 - cap_rect.x1 < page_w * 0.18
            )
            full_width_graphic = graphic.width > page_w * 0.55
            same_row = (
                (same_half or adjacent_same_band)
                and
                vertical_overlap > -20
                and abs(graphic.x0 - cap_rect.x1) < page_w * 0.35
                and graphic.x1 > cap_rect.x1
            )
            caption_inside_graphic = (
                caption_overlap_ratio > 0.35
                and horizontal_overlap > min(cap_rect.width, graphic.width) * 0.45
            )
            above_caption = (
                (same_half or full_width_graphic)
                and
                graphic.y1 <= cap_rect.y0 + 18
                and cap_rect.y0 - graphic.y1 < page_h * 0.50
                and horizontal_overlap > -page_w * 0.12
            )
            below_caption = (
                (same_half or full_width_graphic)
                and
                graphic.y0 >= cap_rect.y1 - 12
                and graphic.y0 - cap_rect.y1 < page_h * 0.18
                and horizontal_overlap > -page_w * 0.12
            )
            if caption_inside_graphic or same_row or above_caption or below_caption:
                if caption_inside_graphic or same_row:
                    same_row_graphics.append(graphic)
                elif above_caption:
                    above_graphics.append((max(0.0, cap_rect.y0 - graphic.y1), graphic))
                else:
                    below_graphics.append((max(0.0, graphic.y0 - cap_rect.y1), graphic))

        nearby: list[fitz.Rect] = []
        if same_row_graphics:
            nearby = same_row_graphics
        elif above_graphics:
            if is_table_caption:
                min_gap = min(gap for gap, _ in above_graphics)
                nearby = [graphic for gap, graphic in above_graphics if gap <= min_gap + 60]
            else:
                nearby = _connected_above_graphics(above_graphics)
        elif below_graphics:
            min_gap = min(gap for gap, _ in below_graphics)
            nearby = [graphic for gap, graphic in below_graphics if gap <= min_gap + 60]

        table_text_rects: list[fitz.Rect] = []
        if not nearby:
            if not is_table_caption:
                continue
            for block in text_blocks:
                block_rect = fitz.Rect(block[0], block[1], block[2], block[3])
                if block_rect.y0 < cap_rect.y1 - 2:
                    continue
                if block_rect.y0 - cap_rect.y1 > page_h * 0.18:
                    continue
                horizontal_overlap = min(cap_rect.x1, block_rect.x1) - max(cap_rect.x0, block_rect.x0)
                block_center_x = (block_rect.x0 + block_rect.x1) / 2
                cap_center_x = (cap_rect.x0 + cap_rect.x1) / 2
                same_half = (
                    block_center_x < page_w * 0.52
                    if cap_center_x < page_w * 0.5
                    else block_center_x > page_w * 0.48
                )
                same_column = same_half and horizontal_overlap > -page_w * 0.04
                text = " ".join(_normalize_text(str(block[4])).strip().split())
                compact_table_text = len(text) <= 120
                if same_column and compact_table_text:
                    table_text_rects.append(block_rect)
            if not table_text_rects:
                continue

        visual_rect = _union_rect(nearby + table_text_rects + [cap_rect])
        if visual_rect is None:
            continue

        related_text = []
        if nearby and not is_table_caption:
            include_zone = _expanded_clip(visual_rect, page_rect, margin=14)
            visual_center_x = (visual_rect.x0 + visual_rect.x1) / 2
            visual_spans_columns = visual_rect.width > page_w * 0.55
            label_x0 = max(visual_rect.x0 - page_w * 0.08, page_rect.x0)
            label_x1 = min(visual_rect.x1 + page_w * 0.25, page_rect.x1)
            if not visual_spans_columns:
                if visual_center_x < page_w * 0.5:
                    label_x1 = min(label_x1, page_w * 0.52)
                else:
                    label_x0 = max(label_x0, page_w * 0.48)
            label_zone = fitz.Rect(
                label_x0,
                max(visual_rect.y0 - 14, page_rect.y0),
                label_x1,
                min(visual_rect.y1 + 14, page_rect.y1),
            )
            broad_zone = _expanded_clip(visual_rect, page_rect, margin=40)
            for block in text_blocks:
                block_rect = fitz.Rect(block[0], block[1], block[2], block[3])
                text = str(block[4]).strip()
                if _rect_overlap_ratio(block_rect, cap_rect) > 0.9:
                    related_text.append(block_rect)
                elif len(text) < 50 and _rect_has_meaningful_overlap(block_rect, broad_zone):
                    # Short labels (axis labels, legend keys, callouts) inside the figure
                    related_text.append(block_rect)
                elif (_rect_has_meaningful_overlap(block_rect, include_zone) or _rect_has_meaningful_overlap(block_rect, label_zone)) and (
                    len(text) < 90
                    or (len(text) < 120 and block_rect.get_area() < page_rect.get_area() * 0.02)
                ):
                    related_text.append(block_rect)

        table_text_only_caption = bool(is_table_caption and not nearby)
        final_rect = _union_rect(nearby + table_text_rects + [cap_rect] + related_text)
        if final_rect is None:
            continue
        if table_text_only_caption:
            # For table captions without any graphic anchor, text blocks often cover
            # only header rows. Expand downward to avoid cropping just the table title.
            min_table_h = page_h * 0.12
            if final_rect.height < min_table_h:
                target_bottom = min(page_rect.y1 - 6, final_rect.y0 + min_table_h)
                final_rect = fitz.Rect(final_rect.x0, max(page_rect.y0 + 4, final_rect.y0 - 4), final_rect.x1, target_bottom)
        final_rect = _expanded_clip(final_rect, page_rect, margin=5)

        if any(_rect_overlap_ratio(final_rect, prev) > 0.7 for prev in exclude):
            continue

        image_id = f"visual_p{page_num}_v{i}.png"
        _job_image_path(job_id, image_id).parent.mkdir(parents=True, exist_ok=True)
        mat = _native_image_matrix_for_clip(final_rect)
        page.get_pixmap(matrix=mat, clip=final_rect, alpha=False).save(str(_job_image_path(job_id, image_id)))
        infos.append({
            "image_id": image_id,
            "y_frac": final_rect.y0 / page_h,
            "order_key": _visual_order_key(final_rect, page, text_blocks),
            "source_idx": _estimate_source_idx_from_text_blocks(page, final_rect, text_blocks),
        })
        exclude.append(final_rect)

    return infos, exclude


def _extract_uncaptioned_image_blocks(
    page: fitz.Page, job_id: str, page_num: int, existing_rects: list | None = None
) -> tuple[list[dict], list]:
    """Fallback for sizeable raster images that have no detected Fig./Table caption."""
    if not job_id:
        return [], []

    existing_rects = existing_rects or []
    all_img_rects: list[fitz.Rect] = []
    seen_xrefs: set[int] = set()
    for image in page.get_images(full=True):
        xref = int(image[0])
        if xref in seen_xrefs:
            continue
        seen_xrefs.add(xref)
        try:
            all_img_rects.extend(fitz.Rect(r) for r in page.get_image_rects(xref))
        except Exception:
            continue

    page_area = page.rect.get_area()
    page_rect = page.rect
    min_area = page_area * 0.01
    max_area = page_area * 0.93  # skip full-page backgrounds (scanned page images); allow large inline photos
    full_page_rasters = [r for r in all_img_rects if r.get_area() >= max_area]
    img_rects = [
        r for r in all_img_rects
        if min_area <= r.get_area() < max_area
        and not any(_rect_overlap_ratio(r, er) > 0.5 for er in existing_rects)
    ]
    # Some "native" PDFs are hybrid pages: one full-page raster background with
    # selectable text overlay. In that case there are no inline image rects below
    # `max_area`, so run render-based visual detection as a fallback.
    detected_from_full_page_raster = False
    fallback_words: list[dict] = []
    if not img_rects and full_page_rasters:
        detected_from_full_page_raster = True
        words = _extract_native_page_words(page)
        fallback_words = words
        if words:
            render_mat = fitz.Matrix(2, 2)
            candidate_rects: list[fitz.Rect] = []
            for detector_rects in (
                _scan_page_visual_envelope(page, words, render_mat),
                _scan_layout_visual_rects(page, words, render_mat),
                _scan_density_visual_rects(page, words, render_mat),
                _scan_visual_rects_from_render(page, words, render_mat),
            ):
                for rect in detector_rects:
                    if rect.get_area() < page_area * 0.03:
                        continue
                    if any(_rect_overlap_ratio(rect, er) > 0.5 for er in existing_rects):
                        continue
                    candidate_rects.append(rect)

            candidate_rects.sort(key=lambda r: (r.y0, r.x0))
            merged_candidates: list[fitz.Rect] = []
            for rect in candidate_rects:
                overlap_idx = next(
                    (idx for idx, existing in enumerate(merged_candidates)
                     if _rect_overlap_min_ratio(rect, existing) > 0.55 or _rect_has_meaningful_overlap(rect, existing)),
                    None,
                )
                if overlap_idx is None:
                    merged_candidates.append(fitz.Rect(rect))
                    continue
                existing = merged_candidates[overlap_idx]
                if rect.get_area() > existing.get_area():
                    merged_candidates[overlap_idx] = fitz.Rect(rect)

            expanded_candidates: list[fitz.Rect] = []
            text_blocks = [
                b for b in page.get_text("blocks")
                if len(b) >= 7 and b[6] == 0 and str(b[4]).strip()
            ]
            for rect in merged_candidates:
                # Render-based detectors on hybrid full-page raster PDFs can crop
                # figure labels near the edge. Keep padding tight: these rects are
                # also used to exclude text spans, so wide padding deletes body text.
                rect_center_x = (rect.x0 + rect.x1) / 2
                wide_candidate = rect.width > page_rect.width * 0.55
                right_side_candidate = rect_center_x > page_rect.width * 0.55
                if wide_candidate:
                    pad_left = min(page_rect.width * 0.07, max(page_rect.width * 0.010, rect.width * 0.10))
                    pad_right = min(page_rect.width * 0.07, max(page_rect.width * 0.010, rect.width * 0.10))
                    pad_top = min(page_rect.height * 0.025, max(page_rect.height * 0.005, rect.height * 0.06))
                    pad_bottom = min(page_rect.height * 0.055, max(page_rect.height * 0.008, rect.height * 0.16))
                else:
                    # Right-column figures are often adjacent to body text on the left;
                    # keep left padding conservative to avoid swallowing prose.
                    pad_left_ratio = 0.10 if right_side_candidate else 0.18
                    pad_left_cap = 0.04 if right_side_candidate else 0.08
                    pad_left = min(page_rect.width * pad_left_cap, max(page_rect.width * 0.018, rect.width * pad_left_ratio))
                    pad_right_ratio = 0.08 if right_side_candidate else 0.04
                    pad_right_cap = 0.06 if right_side_candidate else 0.035
                    pad_right_min = 0.010 if right_side_candidate else 0.005
                    pad_right = min(page_rect.width * pad_right_cap, max(page_rect.width * pad_right_min, rect.width * pad_right_ratio))
                    pad_top = min(page_rect.height * 0.040, max(page_rect.height * 0.006, rect.height * 0.10))
                    pad_bottom = min(page_rect.height * 0.080, max(page_rect.height * 0.012, rect.height * 0.30))
                expanded = fitz.Rect(
                    max(page_rect.x0, rect.x0 - pad_left),
                    max(page_rect.y0, rect.y0 - pad_top),
                    min(page_rect.x1, rect.x1 + pad_right),
                    min(page_rect.y1, rect.y1 + pad_bottom),
                )
                for block in text_blocks:
                    block_rect = fitz.Rect(block[0], block[1], block[2], block[3])
                    text = _normalize_text(str(block[4])).strip()
                    body_text = len(text) >= 120
                    vertical_overlap = min(expanded.y1, block_rect.y1) - max(expanded.y0, block_rect.y0)
                    if body_text and vertical_overlap > min(expanded.height, block_rect.height) * 0.18:
                        if block_rect.x0 > rect.x1 - 1 and block_rect.x0 < expanded.x1:
                            expanded.x1 = min(expanded.x1, max(rect.x1 + 2, block_rect.x0 - 3))
                        if block_rect.x1 < rect.x0 + 1 and block_rect.x1 > expanded.x0:
                            expanded.x0 = max(expanded.x0, min(rect.x0 - 2, block_rect.x1 + 3))
                        if (block_rect.x0 < rect.x0
                                and block_rect.x1 > rect.x0
                                and block_rect.x1 < rect.x0 + rect.width * 0.25):
                            expanded.x0 = max(expanded.x0, block_rect.x1 + 3)
                    horizontal_overlap = min(rect.x1, block_rect.x1) - max(rect.x0, block_rect.x0)
                    if horizontal_overlap < min(rect.width, block_rect.width) * 0.18:
                        continue
                    body_text_above_graphic = (
                        body_text
                        and
                        block_rect.y0 < rect.y0
                        and block_rect.y1 <= rect.y0 + page_rect.height * 0.045
                        and block_rect.y1 > expanded.y0
                    )
                    header_above_graphic = (
                        rect.y0 < page_rect.height * 0.18
                        and block_rect.y0 < rect.y0
                        and block_rect.y1 < page_rect.height * 0.13
                        and block_rect.y1 > expanded.y0
                    )
                    if body_text_above_graphic or header_above_graphic:
                        top_limit = rect.y0 + page_rect.height * 0.045
                        expanded.y0 = max(expanded.y0, min(block_rect.y1 + 3, top_limit))
                if expanded.width <= 0 or expanded.height <= 0:
                    continue
                expanded_candidates.append(expanded)

            filtered_rects: list[fitz.Rect] = []
            for r in expanded_candidates:
                if r.width <= page.rect.width * 0.18:
                    continue
                if r.height <= page.rect.height * 0.10:
                    continue
                area_ratio = r.get_area() / max(page_area, 1.0)
                if area_ratio >= 0.90:
                    continue
                # Extremely large candidates on hybrid pages are often near full-page
                # backgrounds. Reject them when they still contain many prose lines.
                if (
                    detected_from_full_page_raster
                    and area_ratio >= 0.70
                    and fallback_words
                    and _count_body_text_lines_in_rect(r, page, fallback_words) >= 4
                ):
                    continue
                filtered_rects.append(r)
            img_rects = filtered_rects

    if not img_rects:
        return [], []

    page_h = max(page_rect.height, 1)
    infos: list[dict] = []
    exclude: list = []

    _tb_for_order = [
        b for b in page.get_text("blocks")
        if len(b) >= 7 and b[6] == 0 and str(b[4]).strip()
    ]
    _two_col_gutter = None
    if _page_has_side_by_side_columns(page, _tb_for_order):
        narrow_tb = [b for b in _tb_for_order if (b[2] - b[0]) <= page_rect.width * 0.55]
        left_tb = [b for b in narrow_tb if ((b[0] + b[2]) / 2) < page_rect.width * 0.5]
        right_tb = [b for b in narrow_tb if ((b[0] + b[2]) / 2) >= page_rect.width * 0.5]
        if left_tb and right_tb:
            _two_col_gutter = (max(b[2] for b in left_tb), min(b[0] for b in right_tb))

    for i, rect in enumerate(img_rects):
        render_rect = fitz.Rect(rect)
        if detected_from_full_page_raster:
            # Hybrid full-page raster PDFs tend to get overly tight candidate boxes.
            # Expand the render clip a bit so figure edges are not cut off.
            pad_x = max(page_rect.width * 0.014, render_rect.width * 0.06)
            pad_top = max(page_rect.height * 0.008, render_rect.height * 0.025)
            pad_bottom = max(page_rect.height * 0.018, render_rect.height * 0.055)
            render_rect = fitz.Rect(
                max(page_rect.x0, render_rect.x0 - pad_x),
                max(page_rect.y0, render_rect.y0 - pad_top),
                min(page_rect.x1, render_rect.x1 + pad_x),
                min(page_rect.y1, render_rect.y1 + pad_bottom),
            )
        clip_margin = 2 if detected_from_full_page_raster else 4
        clip = _expanded_clip(render_rect, page_rect, margin=clip_margin)
        mat = _native_image_matrix_for_clip(clip)
        pix = page.get_pixmap(matrix=mat, clip=clip, alpha=False)
        image_id = f"image_p{page_num}_i{i}.png"
        _job_image_path(job_id, image_id).parent.mkdir(parents=True, exist_ok=True)
        pix.save(str(_job_image_path(job_id, image_id)))
        infos.append({
            "image_id": image_id,
            "y_frac": rect.y0 / page_h,
            "order_key": _visual_order_key(rect, page, _tb_for_order),
            "source_idx": _estimate_source_idx_from_text_blocks(page, rect, _tb_for_order),
        })
        if detected_from_full_page_raster:
            # Hybrid native PDFs already have selectable body text over a full-page
            # raster. Extra exclusion margins fragment paragraphs around figures.
            exclude.append(_expanded_clip(rect, page_rect, margin=3))
        else:
            margin = min(max(rect.height * 0.25, 20), 50)
            ex = _expanded_clip(rect, page_rect, margin=margin)
            if _two_col_gutter is not None and rect.width < page_rect.width * 0.55:
                left_col_x1, right_col_x0 = _two_col_gutter
                rect_cx = (rect.x0 + rect.x1) / 2
                if rect_cx < page_rect.width * 0.5 and rect.x1 < right_col_x0:
                    ex = fitz.Rect(ex.x0, ex.y0, min(ex.x1, right_col_x0 - 3), ex.y1)
                elif rect_cx >= page_rect.width * 0.5 and rect.x0 > left_col_x1:
                    ex = fitz.Rect(max(ex.x0, left_col_x1 + 3), ex.y0, ex.x1, ex.y1)
            exclude.append(ex)

    return infos, exclude


def _is_formula_text(text: str) -> bool:
    raw = text or ""
    clean = " ".join(_normalize_text(text).strip().split())
    has_control_formula_char = any(ch in raw for ch in ("\x01", "\x02", "\x03"))
    if not clean and has_control_formula_char:
        return True
    if not clean:
        return False
    if len(clean) > 140:
        return False
    if _looks_like_reference_text(clean):
        return False
    lower_clean = clean.lower()
    if lower_clean.startswith(("where ", "use ", "can ", "soft ", "hard ")):
        return False
    if re.search(r"[A-Za-z]{6,}", clean) and len(re.findall(r"[A-Za-z]{4,}", clean)) >= 3:
        return False
    if has_control_formula_char:
        return True
    if re.fullmatch(r"\(?\d+[A-Za-z]?\)?", clean):
        return False
    math_symbols = ("\x01", "\x02", "\x03", "α", "β", "γ", "λ", "∑", "√", "∞", "≤", "≥", "±")
    core_math_ops = ("=", "+", "-", "/", "×", "÷", "±", "≤", "≥")
    if any(ch in clean for ch in math_symbols) and any(op in clean for op in core_math_ops):
        return True
    if any(0x1D400 <= ord(ch) <= 0x1D7FF for ch in clean) and any(op in clean for op in core_math_ops):
        return True
    if re.search(r"(^|[\s(])[A-Za-z]{1,4}\s*=", clean):
        return True
    if re.search(r"=\s*[-+A-Za-z0-9(]", clean):
        return True
    if re.search(r"\+\s*[A-Za-z0-9]", clean) and re.search(r"\(\d+\)", clean):
        return True
    if re.search(r"\(\d+\)", clean) and any(op in clean for op in ("+", "/", "=")):
        return True
    return False


def _looks_like_reference_text(text: str) -> bool:
    clean = " ".join((text or "").split())
    if not clean:
        return False
    ref_terms = (
        "J Hydrol",
        "J Irrig",
        "Agric Water",
        "Academic",
        "Trans Jpn",
        "Geological survey",
        "Ministry of",
        "National Agricultural",
        "Association",
        "Research bulletin",
        "Eng ",
    )
    lower = clean.lower()
    if re.search(r"\b(19|20)\d{2}\b", clean) and any(term.lower() in lower for term in ref_terms):
        return True
    if re.search(r"\b[A-Z][A-Za-z-]+ [A-Z]{1,3},", clean) and re.search(r"\b(19|20)\d{2}\b", clean):
        return True
    if re.search(r"\b\d+\(\d+\):\d+", clean):
        return True
    return False


def _is_formula_adjunct_text(text: str) -> bool:
    raw = text or ""
    clean = " ".join(_normalize_text(text).strip().split())
    has_control_formula_char = any(ch in raw for ch in ("\x01", "\x02", "\x03"))
    if not clean and has_control_formula_char:
        return True
    if not clean:
        return False
    if len(clean) > 100:
        return False
    if re.search(r"[A-Za-z]{6,}", clean):
        return False
    if has_control_formula_char and len(clean) <= 70:
        return True
    if re.fullmatch(r"\(?\d+[A-Za-z]?\)?", clean):
        return True
    if re.fullmatch(r"[A-Za-z][A-Za-z0-9]{0,5}(?:net|ref|sat|pot)?", clean):
        return True
    if clean in {"α", "β", "γ", "λ", "l"}:
        return True
    if any(ch in clean for ch in ("=", "+", "-", "/", "×", "÷", "±", "≤", "≥")):
        return True
    if any(0x1D400 <= ord(ch) <= 0x1D7FF for ch in clean) and len(clean) <= 100:
        return True
    if (
        len(clean) <= 70
        and re.search(r"[()+\-??]", clean)
        and re.search(r"[A-Za-z0-9]", clean)
        and not re.search(r"[A-Za-z]{6,}", clean)
    ):
        return True
    return False


def _extract_formula_span_rects(page: fitz.Page, existing_rects: list | None = None) -> list[fitz.Rect]:
    """Find small formulas split into spans inside a text line, such as stacked d2f/dx2."""
    existing_rects = existing_rects or []
    raw = page.get_text("rawdict")
    page_w = max(page.rect.width, 1)
    span_items: list[tuple[fitz.Rect, str, bool]] = []
    core_ops = ("=", "+", "-", "/", "×", "÷", "±", "≤", "≥")
    image_rects: list[fitz.Rect] = []
    seen_xrefs: set[int] = set()
    for image in page.get_images(full=True):
        xref = int(image[0])
        if xref in seen_xrefs:
            continue
        seen_xrefs.add(xref)
        try:
            image_rects.extend(fitz.Rect(r) for r in page.get_image_rects(xref))
        except Exception:
            continue
    image_rects.extend(_collect_graphic_rects(page))
    image_rects = [
        fitz.Rect(
            max(ir.x0 - 20, page.rect.x0),
            max(ir.y0 - 20, page.rect.y0),
            min(ir.x1 + 20, page.rect.x1),
            min(ir.y1 + 90, page.rect.y1),
        )
        for ir in image_rects
    ]

    for block in raw.get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                chars = span.get("chars", [])
                text = "".join(ch.get("c", "") for ch in chars) if chars else span.get("text", "")
                clean = " ".join(_normalize_text(text).strip().split())
                if not clean:
                    continue
                if re.fullmatch(r"[A-Za-z]{1,5}-", clean):
                    # Native PDFs split justified words into tiny hyphenated spans
                    # such as "de-" or "pri-"; stacked spans are prose, not formulas.
                    continue
                rect = fitz.Rect(span.get("bbox", line.get("bbox", block.get("bbox"))))
                if any(_rect_overlap_ratio(rect, er) > 0.45 for er in existing_rects):
                    continue
                if any(_rect_overlap_ratio(rect, ir) > 0.35 for ir in image_rects):
                    continue
                has_math_alpha = any(0x1D400 <= ord(ch) <= 0x1D7FF for ch in clean)
                has_core = any(op in clean for op in core_ops)
                has_long_word = re.search(r"[A-Za-z]{4,}", clean) is not None
                if not has_math_alpha and not has_core:
                    continue
                if has_long_word:
                    continue
                if rect.width > page_w * 0.45:
                    continue
                span_items.append((rect, clean, has_core))

    groups: list[list[tuple[fitz.Rect, str, bool]]] = []
    for item in sorted(span_items, key=lambda it: (it[0].y0, it[0].x0)):
        rect, _, _ = item
        placed = False
        for group in groups:
            group_rect = _union_rect([g[0] for g in group])
            if group_rect is None:
                continue
            vertical_overlap = min(group_rect.y1, rect.y1) - max(group_rect.y0, rect.y0)
            same_stack = abs(((rect.x0 + rect.x1) / 2) - ((group_rect.x0 + group_rect.x1) / 2)) < 18 and vertical_overlap > -18
            if same_stack:
                group.append(item)
                placed = True
                break
        if not placed:
            groups.append([item])

    rects: list[fitz.Rect] = []
    for group in groups:
        group_rect = _union_rect([g[0] for g in group])
        if group_rect is None:
            continue
        multi_part = len(group) >= 2 and group_rect.height > 12 and group_rect.width <= page_w * 0.18
        if multi_part:
            rects.append(group_rect)
    return rects


def _extract_formula_images(
    page: fitz.Page, job_id: str, page_num: int, existing_rects: list | None = None
) -> tuple[list[dict], list, list]:
    """Render isolated equation/formula regions as images and exclude them."""
    if not job_id:
        return [], [], []
    # Hybrid OCR PDFs often contain a full-page raster + noisy text overlay.
    # Formula heuristics on that overlay easily produce tiny false positives.
    if _page_has_fullpage_raster_background(page):
        return [], [], []
    existing_rects = existing_rects or []
    text_blocks = [
        b for b in page.get_text("blocks")
        if len(b) >= 7 and b[6] == 0 and str(b[4]).strip()
    ]
    core_candidates: list[tuple[fitz.Rect, str]] = []
    adjunct_candidates: list[tuple[fitz.Rect, str]] = []
    for block in text_blocks:
        rect = fitz.Rect(block[0], block[1], block[2], block[3])
        if any(_rect_overlap_ratio(rect, er) > 0.45 for er in existing_rects):
            continue
        text = str(block[4]).strip()
        if _is_formula_text(text):
            core_candidates.append((rect, text))
        elif _is_formula_adjunct_text(text):
            adjunct_candidates.append((rect, text))

    if not core_candidates:
        core_candidates = []

    page_rect = page.rect
    page_w = max(page_rect.width, 1)
    page_h = max(page_rect.height, 1)
    has_columns = _page_has_side_by_side_columns(page, text_blocks)
    by_column: dict[int, list[tuple[fitz.Rect, str]]] = {}
    for rect, text in core_candidates:
        center = (rect.x0 + rect.x1) / 2
        col = (0 if center < page_w * 0.5 else 1) if has_columns else 0
        by_column.setdefault(col, []).append((rect, text))

    formula_rects: list[fitz.Rect] = []
    for col_items in by_column.values():
        col_items.sort(key=lambda item: (item[0].y0, item[0].x0))
        current: list[fitz.Rect] = []
        current_bottom = 0.0
        for rect, _ in col_items:
            if not current:
                current = [rect]
                current_bottom = rect.y1
                continue
            gap = rect.y0 - current_bottom
            if gap <= 24:
                current.append(rect)
                current_bottom = max(current_bottom, rect.y1)
            else:
                merged = _union_rect(current)
                if merged is not None:
                    formula_rects.append(merged)
                current = [rect]
                current_bottom = rect.y1
        merged = _union_rect(current)
        if merged is not None:
            formula_rects.append(merged)

    expanded_formula_rects: list[fitz.Rect] = []
    for rect in formula_rects:
        group = [rect]
        rect_center_x = (rect.x0 + rect.x1) / 2
        rect_left_half = rect_center_x < page_w * 0.5
        for adj_rect, _ in adjunct_candidates:
            adj_center_x = (adj_rect.x0 + adj_rect.x1) / 2
            same_half = adj_center_x < page_w * 0.52 if rect_left_half else adj_center_x > page_w * 0.48
            same_band = adj_rect.y0 <= rect.y1 + 18 and adj_rect.y1 >= rect.y0 - 18
            close_x = adj_rect.x1 >= rect.x0 - page_w * 0.12 and adj_rect.x0 <= rect.x1 + page_w * 0.18
            if (same_half or (same_band and close_x)) and same_band and close_x:
                group.append(adj_rect)
        expanded_formula_rects.append(_union_rect(group) or rect)
    formula_rects = []
    for rect in sorted(expanded_formula_rects, key=lambda r: (r.y0, r.x0)):
        merged_into_existing = False
        for existing in formula_rects:
            vertical_overlap = min(existing.y1, rect.y1) - max(existing.y0, rect.y0)
            horizontal_gap = max(0.0, max(rect.x0 - existing.x1, existing.x0 - rect.x1))
            same_column = _same_reading_column(existing, rect, page_w)
            same_formula_band = (
                vertical_overlap > -18
                and same_column
                and horizontal_gap < page_w * 0.08
            )
            if same_formula_band:
                existing.include_rect(rect)
                merged_into_existing = True
                break
        if not merged_into_existing:
            formula_rects.append(fitz.Rect(rect))

    span_formula_rects = _extract_formula_span_rects(page, existing_rects)
    for span_rect in span_formula_rects:
        if span_rect.x1 > page_w * 0.92 and span_rect.width < page_w * 0.06:
            continue
        if any(_rect_overlap_ratio(span_rect, rect) > 0.3 for rect in formula_rects):
            continue
        formula_rects.append(span_rect)

    infos: list[dict] = []
    exclude: list = []       # rendering clips (8 pt margin) — used for dedup + image-block overlap checks
    tight_rects: list = []   # core formula rects (no margin) — used for text-layer exclusion
    for i, rect in enumerate(sorted(formula_rects, key=lambda r: _visual_order_key(r, page, text_blocks))):
        if rect.get_area() < 20:
            continue
        include_zone = _expanded_clip(rect, page_rect, margin=18)
        adjuncts = []
        rect_center_x = (rect.x0 + rect.x1) / 2
        rect_left_half = rect_center_x < page_w * 0.5
        rect_spans_columns = rect.width > page_w * 0.55
        for block in text_blocks:
            block_rect = fitz.Rect(block[0], block[1], block[2], block[3])
            if any(_rect_overlap_ratio(block_rect, er) > 0.45 for er in existing_rects):
                continue
            block_center_x = (block_rect.x0 + block_rect.x1) / 2
            same_half = block_center_x < page_w * 0.52 if rect_left_half else block_center_x > page_w * 0.48
            same_row = block_rect.y0 <= rect.y1 + 14 and block_rect.y1 >= rect.y0 - 14
            same_column = _same_reading_column(rect, block_rect, page_w)
            close_x = block_rect.x1 >= rect.x0 - page_w * 0.08 and block_rect.x0 <= rect.x1 + page_w * 0.08
            if (
                same_column
                and (same_half or rect_spans_columns or (same_row and close_x))
                and (_rect_has_meaningful_overlap(block_rect, include_zone) or same_row)
                and _is_formula_adjunct_text(str(block[4]))
            ):
                adjuncts.append(block_rect)
        full_rect = _union_rect([rect] + adjuncts) or rect
        clip = _expanded_clip(full_rect, page_rect, margin=8)
        if any(_rect_overlap_ratio(clip, prev) > 0.4 for prev in exclude):
            continue
        if any(_rect_overlap_ratio(clip, fitz.Rect(er)) > 0.40 for er in existing_rects):
            continue  # already captured as part of a figure / visual block
        image_id = f"formula_p{page_num}_f{i}.png"
        _job_image_path(job_id, image_id).parent.mkdir(parents=True, exist_ok=True)
        mat = _native_image_matrix_for_clip(clip)
        page.get_pixmap(matrix=mat, clip=clip, alpha=False).save(str(_job_image_path(job_id, image_id)))
        order_key = clip.y0 / page_h if clip.width > page_w * 0.45 else _visual_order_key(clip, page, text_blocks)
        infos.append({
            "image_id": image_id,
            "y_frac": clip.y0 / page_h,
            "order_key": order_key,
            "source_idx": _estimate_source_idx_from_text_blocks(
                page,
                clip,
                text_blocks,
                exclude_rects=[r for r in formula_rects if _rect_overlap_ratio(r, clip) < 0.7],
            ),
        })
        exclude.append(clip)
        tight_rects.append(fitz.Rect(full_rect))  # tight rect for text exclusion

    return infos, exclude, tight_rects


def _extract_table_images(page: fitz.Page, job_id: str, page_num: int, table_rects: list) -> list[dict]:
    """Render each table region as a PNG and save to job assets.
    Returns list of {image_id, y_frac} for later block insertion."""
    if not table_rects or not job_id:
        return []
    page_rect = page.rect
    page_h = max(page_rect.height, 1)
    result = []
    for i, rect in enumerate(table_rects):
        clip = _expanded_clip(fitz.Rect(rect), page_rect, margin=4)
        mat = _native_image_matrix_for_clip(clip)
        pix = page.get_pixmap(matrix=mat, clip=clip, alpha=False)
        image_id = f"table_p{page_num}_t{i}.png"
        img_path = _job_image_path(job_id, image_id)
        pix.save(str(img_path))
        result.append({
            "image_id": image_id,
            "y_frac": rect.y0 / page_h,
            "order_key": _visual_order_key(fitz.Rect(rect), page),
            "source_idx": _estimate_source_idx_from_text_blocks(page, fitz.Rect(rect)),
        })
    return result


def _merge_table_blocks(translated_blocks: list, page_table_blocks: list) -> list:
    """Insert table image blocks into translated_blocks at positions estimated from y_frac."""
    for page_idx, table_infos in enumerate(page_table_blocks or []):
        if page_idx >= len(translated_blocks) or not table_infos:
            continue
        page_blocks = translated_blocks[page_idx]
        n = len(page_blocks)
        large_scan_visual = any(
            (float(tinfo.get("right", 0)) - float(tinfo.get("left", 0))) > 0.65
            and (float(tinfo.get("bottom", 0)) - float(tinfo.get("top", 0))) > 0.65
            for tinfo in table_infos
            if "left" in tinfo and "right" in tinfo and "top" in tinfo and "bottom" in tinfo
        )
        if large_scan_visual:
            page_blocks = [block for block in page_blocks if block.get("type") == "image"]
        for tinfo in table_infos:
            position = tinfo.get("order_key", tinfo.get("y_frac", 0))
            position = max(0.0, min(1.0, float(position)))
            source_span = max(n - 1, 1) * TRANSLATION_BLOCK_STEP
            image_id = str(tinfo.get("image_id", ""))
            offset = TRANSLATION_BLOCK_STEP * 2 if image_id.startswith("formula_") else TRANSLATION_BLOCK_STEP / 4
            y_source_idx = round(position * source_span - offset, 6)
            if image_id.startswith("formula_"):
                source_idx = y_source_idx
            elif "source_idx" in tinfo:
                source_idx = round(float(tinfo["source_idx"]), 6)
                if source_idx < -TRANSLATION_BLOCK_STEP or source_idx > source_span + TRANSLATION_BLOCK_STEP:
                    source_idx = y_source_idx
            else:
                source_idx = y_source_idx
            page_blocks.append({
                "type": "image",
                "source_idx": source_idx,
                "image_id": image_id,
                "name": "",
                "width": 0,
            })
        translated_blocks[page_idx] = sorted(page_blocks, key=lambda b: b.get("source_idx", 0))
    return translated_blocks


def _get_ordered_text_blocks(page: fitz.Page, exclude_rects: list | None = None) -> list[tuple]:
    """Return text blocks in visual reading order (shared logic for text and bbox extraction)."""
    blocks = _filtered_text_blocks(page, exclude_rects)
    text_blocks = [b for b in blocks if b[6] == 0 and b[4].strip()]
    if not text_blocks:
        return []

    page_width = page.rect.width

    if not _page_has_side_by_side_columns(page, text_blocks):
        def _single_column_key(b: tuple) -> tuple:
            center_y = (float(b[1]) + float(b[3])) / 2
            row = round(center_y / 24.0)
            return (row, float(b[0]), float(b[1]))
        return sorted(text_blocks, key=_single_column_key)

    full_width = [b for b in text_blocks if (b[2] - b[0]) > page_width * 0.55]
    narrow = [b for b in text_blocks if (b[2] - b[0]) <= page_width * 0.55]

    if not narrow:
        return sorted(full_width, key=lambda b: b[1])

    x0_sorted = sorted(b[0] for b in narrow)
    col_starts = [x0_sorted[0]]
    for i in range(1, len(x0_sorted)):
        if x0_sorted[i] - x0_sorted[i - 1] > page_width * 0.10:
            col_starts.append(x0_sorted[i])
    num_cols = min(len(col_starts), 4)

    if num_cols <= 1:
        return sorted(text_blocks, key=lambda b: (b[1], b[0]))

    columns: list[list] = [[] for _ in range(num_cols)]
    for b in narrow:
        ci = min(range(num_cols), key=lambda i: abs(b[0] - col_starts[i]))
        columns[ci].append(b)
    for col in columns:
        col.sort(key=lambda b: b[1])

    full_sorted = sorted(full_width, key=lambda b: b[1])
    col_y_min = min(col[0][1] for col in columns if col)
    col_y_max = max(col[-1][3] for col in columns if col)

    top_fw = [b for b in full_sorted if b[3] <= col_y_min]
    mid_fw = [b for b in full_sorted if b[1] < col_y_max and b[3] > col_y_min]
    bot_fw = [b for b in full_sorted if b[1] >= col_y_max]

    ordered = list(top_fw)
    fi = 0
    for col in columns:
        for cb in col:
            while fi < len(mid_fw) and mid_fw[fi][1] < cb[1]:
                if mid_fw[fi] not in ordered:
                    ordered.append(mid_fw[fi])
                fi += 1
            ordered.append(cb)
    while fi < len(mid_fw):
        if mid_fw[fi] not in ordered:
            ordered.append(mid_fw[fi])
        fi += 1
    ordered.extend(bot_fw)
    return ordered


def _extract_text_column_aware(page: fitz.Page, exclude_rects: list | None = None) -> str:
    """Extract page text in visual reading order, handling 1–N column layouts."""
    ordered = _get_ordered_text_blocks(page, exclude_rects)
    return "\n\n".join(b[4].strip() for b in ordered)


def _extract_source_block_positions(page: fitz.Page, exclude_rects: list | None = None) -> list[dict]:
    """Return normalized bbox + char count for each text block in reading order."""
    ordered = _get_ordered_text_blocks(page, exclude_rects)
    W, H = page.rect.width, page.rect.height
    result = []
    for idx, b in enumerate(ordered):
        x0, y0, x1, y1, text = b[0], b[1], b[2], b[3], b[4]
        clean_text = _normalize_text(str(text)).strip()
        if not clean_text:
            continue
        box = _normalize_line_box(x0, y0, x1, y1, W, H)
        if box:
            result.append({
                **box,
                "chars": len(clean_text),
                "source_idx": round(idx * TRANSLATION_BLOCK_STEP, 6),
                "text": clean_text,
            })
    return result


def _extract_text_native(
    pdf_path: Path, job_id: str | None = None
) -> tuple[list[str], list[list[dict]], list[list[dict]], list[list[dict]], list[list[dict]]]:
    doc = fitz.open(str(pdf_path))
    pages = []
    page_lines = []
    page_words = []
    page_source_blocks: list[list[dict]] = []
    page_table_blocks: list[list[dict]] = []
    for page_num, page in enumerate(doc):
        table_rects = _find_table_rects(page)
        visual_infos, visual_rects = _extract_captioned_visuals(page, job_id, page_num, table_rects) if job_id else ([], [])
        formula_infos, formula_clips, formula_tight = _extract_formula_images(page, job_id, page_num, visual_rects) if job_id else ([], [], [])
        image_infos, image_rects = _extract_uncaptioned_image_blocks(page, job_id, page_num, visual_rects + formula_clips) if job_id else ([], [])
        text_exclude_rects: list[fitz.Rect] = []
        for rect in table_rects + visual_rects:
            r = fitz.Rect(rect)
            margin = min(max(min(r.width, r.height) * 0.03, 6.0), 18.0)
            text_exclude_rects.append(_expanded_clip(r, page.rect, margin=margin))
        for rect in formula_tight:
            # Use the tight formula rect (no render-margin) + a small 2 pt buffer so
            # text immediately adjacent to the formula is not incorrectly excluded.
            text_exclude_rects.append(_expanded_clip(fitz.Rect(rect), page.rect, margin=2.0))
        for rect in image_rects:
            # Uncaptioned image extraction already adds the needed safety margin.
            # A second large expansion fragments nearby native-PDF paragraphs.
            text_exclude_rects.append(_expanded_clip(fitz.Rect(rect), page.rect, margin=1.0))
        pages.append(_normalize_text(_extract_text_column_aware(page, exclude_rects=text_exclude_rects)))
        page_lines.append(_extract_native_page_lines(page))
        page_words.append(_extract_native_page_words(page))
        page_source_blocks.append(_extract_source_block_positions(page, exclude_rects=text_exclude_rects))
        uncovered_tables = [
            rect for rect in table_rects
            if not any(_rect_overlap_ratio(fitz.Rect(rect), visual_rect) > 0.5 for visual_rect in visual_rects)
        ]
        table_infos = _extract_table_images(page, job_id, page_num, uncovered_tables) if job_id else []
        page_table_blocks.append(table_infos + visual_infos + formula_infos + image_infos)
    doc.close()
    return pages, page_lines, page_words, page_source_blocks, page_table_blocks


def _extract_ocr_page_words(data: dict, image_width: int, image_height: int) -> list[dict]:
    words = []
    total = len(data.get("text", []))
    conf_list = data.get("conf", [])
    for i in range(total):
        text = (data["text"][i] or "").strip()
        if not text:
            continue
        conf = int(conf_list[i]) if i < len(conf_list) else -1
        if conf < 0:
            continue
        left = int(data["left"][i])
        top = int(data["top"][i])
        width = int(data["width"][i])
        height = int(data["height"][i])
        box = _normalize_line_box(left, top, left + width, top + height, image_width, image_height)
        if box:
            box["text"] = text
            words.append(box)
    return sorted(words, key=lambda w: (w["top"], w["left"]))


def _group_ocr_paragraphs(page_lines: list[dict]) -> list[dict]:
    """Port of buildScannedOverlayParagraphs from frontend JS.

    Works in normalized (0-1) coordinates. Returns
    [{idx, left, right, top, bottom}, ...] in visual reading order.
    """
    clean = [
        l for l in page_lines
        if l.get("right", 0) > l.get("left", 0) and l.get("bottom", 0) > l.get("top", 0)
    ]
    if not clean:
        return []

    mid_page = 0.5
    mids = [(l["left"] + l["right"]) / 2 for l in clean]
    left_count = sum(1 for x in mids if x < mid_page * 0.92)
    right_count = sum(1 for x in mids if x > mid_page * 1.08)
    two_col = (
        left_count >= 5 and right_count >= 5
        and min(left_count, right_count) / len(clean) > 0.22
    )

    sort_key = lambda l: (l["top"], l["left"])  # noqa: E731
    if two_col and len(clean) >= 8:
        full_w = [l for l in clean if (l["right"] - l["left"]) > 0.68]
        narrow = [l for l in clean if (l["right"] - l["left"]) <= 0.68]
        left_lines = [l for l in narrow if (l["left"] + l["right"]) / 2 < mid_page]
        right_lines = [l for l in narrow if (l["left"] + l["right"]) / 2 >= mid_page]
        if narrow:
            min_nt = min(l["top"] for l in narrow)
            max_nb = max(l["bottom"] for l in narrow)
        else:
            min_nt, max_nb = 1.0, 0.0
        top_full = sorted([l for l in full_w if l["bottom"] <= min_nt + 0.01], key=sort_key)
        bot_full = sorted([l for l in full_w if l["top"] > max_nb - 0.01], key=sort_key)
        tf_ids = {id(l) for l in top_full + bot_full}
        mid_full = sorted([l for l in full_w if id(l) not in tf_ids], key=sort_key)
        ordered = (
            top_full
            + sorted(left_lines, key=sort_key)
            + sorted(right_lines, key=sort_key)
            + mid_full
            + bot_full
        )
    else:
        ordered = sorted(clean, key=sort_key)

    if not ordered:
        return []

    heights = sorted([max(l["bottom"] - l["top"], 0.001) for l in ordered])
    median_h = heights[len(heights) // 2] if heights else 0.015

    def line_col(line):
        mid_x = (line["left"] + line["right"]) / 2
        if (line["right"] - line["left"]) > 0.68:
            return "full"
        return "left" if mid_x < mid_page else "right"

    paragraphs: list[dict] = []
    cur: list[dict] = []
    cur_col: str | None = None

    def flush():
        nonlocal cur, cur_col
        if not cur:
            return
        paragraphs.append({
            "left": min(l["left"] for l in cur),
            "right": max(l["right"] for l in cur),
            "top": cur[0]["top"],
            "bottom": cur[-1]["bottom"],
        })
        cur = []
        cur_col = None

    for line in ordered:
        col = line_col(line)
        if not cur:
            cur = [line]
            cur_col = col
            continue
        prev = cur[-1]
        gap = line["top"] - prev["bottom"]
        same_col = col == cur_col or col == "full" or cur_col == "full"
        cur_left = min(l["left"] for l in cur)
        cur_right = max(l["right"] for l in cur)
        cur_width = max(cur_right - cur_left, 0.001)
        prev_width = max(prev["right"] - prev["left"], 0.001)
        line_indent = line["left"] - cur_left
        indent_threshold = max(0.012, median_h * 0.8)
        first_line_indent = (
            same_col
            and len(cur) >= 2
            and line_indent > indent_threshold
            and gap > -median_h * 0.2
        )
        after_short_line = (
            same_col
            and len(cur) >= 2
            and prev_width < cur_width * 0.78
            and line["left"] <= cur_left + indent_threshold
            and gap > -median_h * 0.2
        )
        new_para = not same_col or gap > median_h * 0.85 or first_line_indent or after_short_line
        if new_para:
            flush()
            cur = [line]
            cur_col = col
        else:
            cur.append(line)
    flush()

    result = []
    for p in paragraphs:
        w = p["right"] - p["left"]
        h = p["bottom"] - p["top"]
        if w > 0.04 and h > median_h * 0.7:
            result.append({
                "idx": len(result),
                "left": round(p["left"], 6),
                "right": round(p["right"], 6),
                "top": round(p["top"], 6),
                "bottom": round(p["bottom"], 6),
            })
    return result


def _assign_words_to_paragraphs(page_words: list[dict], paragraphs: list[dict]) -> list[str]:
    """Collect OCR word text for each paragraph bounding box.

    Returns a list of strings (one per paragraph) assembled from words whose
    vertical centre falls inside the paragraph's vertical range.
    """
    if not paragraphs or not page_words:
        return ["" for _ in paragraphs]

    para_words: list[list[dict]] = [[] for _ in paragraphs]
    for word in page_words:
        wcy = (word["top"] + word["bottom"]) / 2
        wcx = (word["left"] + word["right"]) / 2
        for i, para in enumerate(paragraphs):
            if (
                para["top"] - 0.005 <= wcy <= para["bottom"] + 0.005
                and para["left"] - 0.02 <= wcx <= para["right"] + 0.02
            ):
                para_words[i].append(word)
                break  # assign to first matching para only

    result = []
    for pw in para_words:
        pw.sort(key=lambda w: (round(w["top"] * 200), w["left"]))
        result.append(" ".join(w["text"] for w in pw))
    return result


def _extract_ocr_page_lines(data: dict, image_width: int, image_height: int) -> list[dict]:
    grouped = {}
    total = len(data.get("text", []))

    for i in range(total):
        text = (data["text"][i] or "").strip()
        if not text:
            continue

        key = (
            int(data["block_num"][i]),
            int(data["par_num"][i]),
            int(data["line_num"][i]),
        )
        left = int(data["left"][i])
        top = int(data["top"][i])
        width = int(data["width"][i])
        height = int(data["height"][i])

        entry = grouped.setdefault(key, {
            "left": left,
            "top": top,
            "right": left + width,
            "bottom": top + height,
        })
        entry["left"] = min(entry["left"], left)
        entry["top"] = min(entry["top"], top)
        entry["right"] = max(entry["right"], left + width)
        entry["bottom"] = max(entry["bottom"], top + height)

    lines = []
    for entry in grouped.values():
        box = _normalize_line_box(
            entry["left"],
            entry["top"],
            entry["right"],
            entry["bottom"],
            image_width,
            image_height,
        )
        if box:
            lines.append(box)

    return sorted(lines, key=lambda item: (item["top"], item["left"]))


def _extract_text_ocr(pdf_path: Path, lang: str = "chi_tra+chi_sim+eng") -> tuple[list[str], list[list[dict]], list[list[dict]]]:
    doc = fitz.open(str(pdf_path))
    pages = []
    page_lines = []
    page_words = []
    for page in doc:
        mat = fitz.Matrix(300 / 72, 300 / 72)
        pix = page.get_pixmap(matrix=mat)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        try:
            text = pytesseract.image_to_string(img, lang=lang, config="--oem 3 --psm 1")
            data = pytesseract.image_to_data(img, lang=lang, output_type=pytesseract.Output.DICT, config="--oem 3 --psm 1")
        except pytesseract.TesseractNotFoundError:
            raise RuntimeError(
                "找不到 Tesseract OCR，程式無法辨識掃描版 PDF。\n"
                "請確認 ButterLayer 資料夾內有 tesseract\\ 子資料夾，\n"
                "或重新下載完整版 ButterLayer。"
            )
        pages.append(_normalize_text(text.strip()))
        page_lines.append(_extract_ocr_page_lines(data, img.width, img.height))
        page_words.append(_extract_ocr_page_words(data, img.width, img.height))
    doc.close()
    return pages, page_lines, page_words


def _needs_ocr(pages: list[str], threshold: int = 50) -> bool:
    import re as _re
    non_empty = [p for p in pages if len(p) >= threshold]
    if len(non_empty) < max(1, len(pages) * 0.5):
        return True
    def _is_garbled(text: str) -> bool:
        non_ws = [c for c in text if not c.isspace()]
        if len(non_ws) < 100:
            return False
        alpha_r = sum(1 for c in non_ws if c.isalpha()) / len(non_ws)
        if alpha_r >= 0.70:
            return False
        word_chars = sum(len(m) for m in _re.findall(r'[a-zA-Z]{2,}', text))
        word_r = word_chars / len(non_ws)
        return alpha_r < 0.65 and word_r < 0.55
    substantial = [p for p in non_empty if sum(1 for c in p if not c.isspace()) >= 100]
    garbled = [p for p in substantial if _is_garbled(p)]
    return len(garbled) >= max(1, int(len(substantial) * 0.10))


def _scan_visual_rects_from_render(
    page: fitz.Page,
    words: list[dict],
    render_matrix: fitz.Matrix,
) -> list[fitz.Rect]:
    """Find non-text visual regions on a scanned page render.

    OCR gap detection misses side-by-side layouts because text can share the
    same vertical band as a figure. This render pass masks OCR words, then
    groups the remaining non-white pixels into candidate visual regions.
    """
    page_rect = page.rect
    page_w = max(page_rect.width, 1)
    page_h = max(page_rect.height, 1)

    try:
        pix = page.get_pixmap(matrix=render_matrix, alpha=False)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    except Exception:
        return []

    if img.width <= 0 or img.height <= 0:
        return []

    scale = min(1.0, 900 / max(img.width, img.height))
    scan_w = max(1, int(img.width * scale))
    scan_h = max(1, int(img.height * scale))
    scan = img.resize((scan_w, scan_h), Image.Resampling.BILINEAR) if scale < 1 else img

    text_mask = Image.new("L", (scan_w, scan_h), 0)
    text_draw = ImageDraw.Draw(text_mask)
    for word in words:
        try:
            left = float(word["left"]) * scan_w
            right = float(word["right"]) * scan_w
            top = float(word["top"]) * scan_h
            bottom = float(word["bottom"]) * scan_h
        except (KeyError, TypeError, ValueError):
            continue
        if right <= left or bottom <= top:
            continue
        pad_x = max(2, (right - left) * 0.18)
        pad_y = max(2, (bottom - top) * 0.35)
        text_draw.rectangle(
            [
                max(0, left - pad_x),
                max(0, top - pad_y),
                min(scan_w, right + pad_x),
                min(scan_h, bottom + pad_y),
            ],
            fill=255,
        )

    gray = scan.convert("L")
    rgb = scan.load()
    gray_px = gray.load()
    text_px = text_mask.load()
    candidate = Image.new("L", (scan_w, scan_h), 0)
    cand_px = candidate.load()

    for y in range(scan_h):
        for x in range(scan_w):
            if text_px[x, y]:
                continue
            r, g, b = rgb[x, y]
            lum = gray_px[x, y]
            saturation = max(r, g, b) - min(r, g, b)
            if lum < 232 or saturation > 24:
                cand_px[x, y] = 255

    candidate = candidate.filter(ImageFilter.MaxFilter(17))
    cand_px = candidate.load()
    text_px = text_mask.load()
    total_area = scan_w * scan_h
    min_area = max(80, int(total_area * 0.0012))
    max_area = int(total_area * 0.82)
    min_w = max(14, int(scan_w * 0.035))
    min_h = max(14, int(scan_h * 0.025))

    visited = bytearray(total_area)
    rects: list[fitz.Rect] = []
    for start_y in range(scan_h):
        row = start_y * scan_w
        for start_x in range(scan_w):
            idx = row + start_x
            if visited[idx] or not cand_px[start_x, start_y]:
                continue
            visited[idx] = 1
            stack = [(start_x, start_y)]
            x0 = x1 = start_x
            y0 = y1 = start_y
            area = 0
            text_hits = 0
            while stack:
                x, y = stack.pop()
                area += 1
                if text_px[x, y]:
                    text_hits += 1
                if x < x0:
                    x0 = x
                elif x > x1:
                    x1 = x
                if y < y0:
                    y0 = y
                elif y > y1:
                    y1 = y

                for nx, ny in ((x - 1, y), (x + 1, y), (x, y - 1), (x, y + 1)):
                    if nx < 0 or nx >= scan_w or ny < 0 or ny >= scan_h:
                        continue
                    nidx = ny * scan_w + nx
                    if visited[nidx] or not cand_px[nx, ny]:
                        continue
                    visited[nidx] = 1
                    stack.append((nx, ny))

            bw = x1 - x0 + 1
            bh = y1 - y0 + 1
            bbox_area = max(bw * bh, 1)
            if area < min_area or area > max_area:
                continue
            if bw < min_w or bh < min_h:
                continue
            if area / bbox_area < 0.025:
                continue
            if text_hits / max(area, 1) > 0.25:
                continue

            margin = max(3, round(min(scan_w, scan_h) * 0.006))
            rx0 = max(0, x0 - margin) / scan_w * page_w
            ry0 = max(0, y0 - margin) / scan_h * page_h
            rx1 = min(scan_w, x1 + margin) / scan_w * page_w
            ry1 = min(scan_h, y1 + margin) / scan_h * page_h
            rect = fitz.Rect(rx0, ry0, rx1, ry1)
            if rect.get_area() < page_rect.get_area() * 0.002:
                continue
            if rect.get_area() > page_rect.get_area() * 0.82:
                continue
            if rect.height > page_h * 0.75 and rect.width < page_w * 0.32:
                continue
            if rect.height > page_h * 0.70 and rect.width < page_w * 0.18:
                continue
            if rect.width < page_w * 0.10 and rect.height > page_h * 0.35:
                continue
            rects.append(rect)

    rects.sort(key=lambda r: (r.y0, r.x0))
    merged: list[fitz.Rect] = []
    for rect in rects:
        joined = False
        for existing in merged:
            # Two regions that share Y-range but are separated in X are side-by-side
            # photos — don't merge them into one wide composite image.
            y_overlap = min(rect.y1, existing.y1) - max(rect.y0, existing.y0)
            x_gap = max(rect.x0 - existing.x1, existing.x0 - rect.x1, 0)
            side_by_side = y_overlap > min(rect.height, existing.height) * 0.3 and x_gap > page_w * 0.01
            if side_by_side:
                continue
            close = (
                rect.x0 <= existing.x1 + page_w * 0.03
                and rect.x1 >= existing.x0 - page_w * 0.03
                and rect.y0 <= existing.y1 + page_h * 0.025
                and rect.y1 >= existing.y0 - page_h * 0.025
            )
            if close or _rect_has_meaningful_overlap(rect, existing):
                existing.include_rect(rect)
                joined = True
                break
        if not joined:
            merged.append(fitz.Rect(rect))

    return [
        rect for rect in merged
        if rect.width >= page_w * 0.055 and rect.height >= page_h * 0.035
        and rect.width >= page_w * 0.10 and rect.height >= page_h * 0.055
        and not (rect.height > page_h * 0.75 and rect.width < page_w * 0.32)
        and not (rect.height > page_h * 0.70 and rect.width < page_w * 0.18)
        and not (rect.width < page_w * 0.10 and rect.height > page_h * 0.35)
    ]


def _ocr_word_lines(words: list[dict]) -> list[dict]:
    rows: list[dict] = []
    for word in sorted(words, key=lambda item: (float(item.get("top", 0)), float(item.get("left", 0)))):
        text = str(word.get("text", "")).strip()
        if not text:
            continue
        try:
            left = float(word["left"])
            right = float(word["right"])
            top = float(word["top"])
            bottom = float(word["bottom"])
        except (KeyError, TypeError, ValueError):
            continue
        if right <= left or bottom <= top:
            continue
        mid_y = (top + bottom) / 2
        target = None
        for row in rows:
            row_h = max(row["bottom"] - row["top"], 0.01)
            if abs(mid_y - row["mid_y"]) <= row_h * 0.75:
                target = row
                break
        if target is None:
            rows.append({
                "left": left,
                "right": right,
                "top": top,
                "bottom": bottom,
                "mid_y": mid_y,
                "words": [text],
            })
        else:
            target["left"] = min(target["left"], left)
            target["right"] = max(target["right"], right)
            target["top"] = min(target["top"], top)
            target["bottom"] = max(target["bottom"], bottom)
            target["mid_y"] = (target["top"] + target["bottom"]) / 2
            target["words"].append(text)
    for row in rows:
        row["text"] = " ".join(row["words"])
    return sorted(rows, key=lambda item: (item["top"], item["left"]))


def _ocr_caption_visual_rects(page: fitz.Page, words: list[dict]) -> list[fitz.Rect]:
    """Infer scanned figure regions from OCR captions such as Fig./Figure/Plate."""
    lines = _ocr_word_lines(words)
    if not lines:
        return []

    page_rect = page.rect
    page_w = max(page_rect.width, 1)
    page_h = max(page_rect.height, 1)
    caption_re = re.compile(r"\b(?:fig(?:ure)?|plate)\s*[\.:]?\s*\d*", re.IGNORECASE)
    rects: list[fitz.Rect] = []

    for idx, line in enumerate(lines):
        caption_text = str(line.get("text", ""))
        if not caption_re.search(caption_text):
            continue
        cap_top = float(line["top"])
        cap_left = float(line["left"])
        cap_right = float(line["right"])
        cap_mid = (cap_left + cap_right) / 2
        if cap_top < 0.08:
            continue

        if cap_right - cap_left > 0.42:
            x0_frac, x1_frac = 0.05, 0.95
        elif cap_mid < 0.50:
            x0_frac, x1_frac = 0.05, 0.52
        else:
            x0_frac, x1_frac = 0.48, 0.95

        above = [
            prev for prev in lines[:idx]
            if float(prev["bottom"]) < cap_top - 0.01
            and not caption_re.search(str(prev.get("text", "")))
            and float(prev["right"]) > x0_frac
            and float(prev["left"]) < x1_frac
        ]
        y0_frac = 0.045
        if above:
            previous = max(above, key=lambda item: float(item["bottom"]))
            gap = cap_top - float(previous["bottom"])
            if gap < 0.05:
                # Caption is embedded in running text; not enough room for a visual above.
                continue
            y0_frac = max(float(previous["bottom"]) + 0.006, cap_top - 0.36)
        else:
            y0_frac = max(0.045, cap_top - 0.42)

        y1_frac = cap_top - 0.006
        if y1_frac - y0_frac < 0.045:
            continue

        rect = fitz.Rect(
            x0_frac * page_w,
            y0_frac * page_h,
            x1_frac * page_w,
            y1_frac * page_h,
        )
        if rect.width < page_w * 0.15 or rect.height < page_h * 0.045:
            continue
        if any(_rect_overlap_ratio(rect, existing) > 0.55 for existing in rects):
            continue
        rects.append(rect)

    return rects


def _scan_density_visual_rects(
    page: fitz.Page,
    words: list[dict],
    render_matrix: fitz.Matrix,
) -> list[fitz.Rect]:
    """Find scanned visuals as dense non-text bands, useful for line drawings."""
    page_rect = page.rect
    page_w = max(page_rect.width, 1)
    page_h = max(page_rect.height, 1)
    try:
        pix = page.get_pixmap(matrix=render_matrix, alpha=False)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    except Exception:
        return []

    scale = min(1.0, 900 / max(img.width, img.height))
    scan_w = max(1, int(img.width * scale))
    scan_h = max(1, int(img.height * scale))
    scan = img.resize((scan_w, scan_h), Image.Resampling.BILINEAR) if scale < 1 else img
    gray = scan.convert("L")
    rgb = scan.load()
    gray_px = gray.load()

    text_mask = Image.new("L", (scan_w, scan_h), 0)
    text_draw = ImageDraw.Draw(text_mask)
    for word in words:
        try:
            left = float(word["left"]) * scan_w
            right = float(word["right"]) * scan_w
            top = float(word["top"]) * scan_h
            bottom = float(word["bottom"]) * scan_h
        except (KeyError, TypeError, ValueError):
            continue
        if right <= left or bottom <= top:
            continue
        pad_x = max(2, (right - left) * 0.12)
        pad_y = max(2, (bottom - top) * 0.25)
        text_draw.rectangle(
            [max(0, left - pad_x), max(0, top - pad_y), min(scan_w, right + pad_x), min(scan_h, bottom + pad_y)],
            fill=255,
        )
    text_px = text_mask.load()

    fg = Image.new("L", (scan_w, scan_h), 0)
    fg_px = fg.load()
    for y in range(scan_h):
        for x in range(scan_w):
            if text_px[x, y]:
                continue
            r, g, b = rgb[x, y]
            lum = gray_px[x, y]
            sat = max(r, g, b) - min(r, g, b)
            if lum < 225 or sat > 24:
                fg_px[x, y] = 255

    rects: list[fitz.Rect] = []
    zones = [(0.06, 0.94), (0.06, 0.52), (0.48, 0.94)]
    min_band_h = max(24, int(scan_h * 0.045))
    for zx0, zx1 in zones:
        x0 = int(scan_w * zx0)
        x1 = int(scan_w * zx1)
        if x1 <= x0:
            continue
        zone_w = x1 - x0
        row_density: list[float] = []
        for y in range(scan_h):
            count = 0
            for x in range(x0, x1):
                if fg_px[x, y]:
                    count += 1
            row_density.append(count / max(zone_w, 1))

        active = [value > 0.045 for value in row_density]
        # Bridge small gaps so sparse architectural line drawings form one band.
        bridge = max(4, int(scan_h * 0.008))
        i = 0
        while i < len(active):
            if active[i]:
                i += 1
                continue
            j = i
            while j < len(active) and not active[j]:
                j += 1
            if i > 0 and j < len(active) and (j - i) <= bridge:
                for k in range(i, j):
                    active[k] = True
            i = j

        y = 0
        while y < scan_h:
            if not active[y]:
                y += 1
                continue
            y0 = y
            while y < scan_h and active[y]:
                y += 1
            y1 = y
            if y1 - y0 < min_band_h:
                continue

            xs: list[int] = []
            for x in range(x0, x1):
                col_count = 0
                for yy in range(y0, y1):
                    if fg_px[x, yy]:
                        col_count += 1
                if col_count / max(y1 - y0, 1) > 0.012:
                    xs.append(x)
            if not xs:
                continue
            bx0 = max(0, min(xs) - 6)
            bx1 = min(scan_w, max(xs) + 6)
            by0 = max(0, y0 - 6)
            by1 = min(scan_h, y1 + 6)
            rect = fitz.Rect(
                bx0 / scan_w * page_w,
                by0 / scan_h * page_h,
                bx1 / scan_w * page_w,
                by1 / scan_h * page_h,
            )
            if rect.width < page_w * 0.14 or rect.height < page_h * 0.045:
                continue
            if rect.y0 < page_h * 0.06 and rect.height < page_h * 0.16:
                continue
            if rect.y0 > page_h * 0.86 and rect.height < page_h * 0.12:
                continue
            if rect.height > page_h * 0.75 and rect.width < page_w * 0.32:
                continue
            if rect.height > page_h * 0.70 and rect.width < page_w * 0.20:
                continue
            word_centers = 0
            rx0 = rect.x0 / page_w
            rx1 = rect.x1 / page_w
            ry0 = rect.y0 / page_h
            ry1 = rect.y1 / page_h
            for word in words:
                try:
                    cx = (float(word["left"]) + float(word["right"])) / 2
                    cy = (float(word["top"]) + float(word["bottom"])) / 2
                except (KeyError, TypeError, ValueError):
                    continue
                if rx0 <= cx <= rx1 and ry0 <= cy <= ry1:
                    word_centers += 1
            area_frac = rect.get_area() / max(page_rect.get_area(), 1)
            if word_centers > max(4, int(area_frac * 120)):
                continue
            if any(_rect_overlap_ratio(rect, existing) > 0.55 for existing in rects):
                continue
            rects.append(rect)

    rects.sort(key=lambda r: (r.y0, r.x0))
    merged: list[fitz.Rect] = []
    for rect in rects:
        for existing in merged:
            if _rect_has_meaningful_overlap(rect, existing):
                existing.include_rect(rect)
                break
        else:
            merged.append(fitz.Rect(rect))
    return merged


def _count_body_text_lines_in_rect(rect: fitz.Rect, page: fitz.Page, words: list[dict]) -> int:
    """Count OCR lines in rect that look like running prose, not map/photo labels."""
    page_w = max(page.rect.width, 1)
    page_h = max(page.rect.height, 1)
    rx0 = rect.x0 / page_w
    rx1 = rect.x1 / page_w
    ry0 = rect.y0 / page_h
    ry1 = rect.y1 / page_h
    count = 0
    for line in _ocr_word_lines(words):
        try:
            left = float(line["left"])
            right = float(line["right"])
            top = float(line["top"])
            bottom = float(line["bottom"])
        except (KeyError, TypeError, ValueError):
            continue
        cx = (left + right) / 2
        cy = (top + bottom) / 2
        if not (rx0 <= cx <= rx1 and ry0 <= cy <= ry1):
            continue
        text = str(line.get("text", "")).strip()
        if not text:
            continue
        width = right - left
        latin_words = re.findall(r"[A-Za-z]{2,}", text)
        cjk_chars = sum(1 for ch in text if 0x4E00 <= ord(ch) <= 0x9FFF)
        looks_like_prose = (
            width > 0.24
            and (
                len(latin_words) >= 5
                or len(text) >= 42
                or cjk_chars >= 16
            )
        )
        if looks_like_prose:
            count += 1
    return count


def _scan_layout_visual_rects(
    page: fitz.Page,
    words: list[dict],
    render_matrix: fitz.Matrix,
) -> list[fitz.Rect]:
    """Find large scanned illustrations whose OCR labels break simple gap detection.

    Maps and architectural plans are often sparse line drawings with many small
    OCR labels.  Density-only detection crops them to the darkest sub-region.
    This pass finds broad foreground bands, then uses OCR line clusters and the
    large gap before body text to keep the whole visual but not the following
    prose.
    """
    page_rect = page.rect
    page_w = max(page_rect.width, 1)
    page_h = max(page_rect.height, 1)
    try:
        pix = page.get_pixmap(matrix=render_matrix, alpha=False)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    except Exception:
        return []

    scale = min(1.0, 900 / max(img.width, img.height))
    scan_w = max(1, int(img.width * scale))
    scan_h = max(1, int(img.height * scale))
    scan = img.resize((scan_w, scan_h), Image.Resampling.BILINEAR) if scale < 1 else img
    gray = scan.convert("L")
    gray_px = gray.load()

    text_mask = Image.new("L", (scan_w, scan_h), 0)
    text_draw = ImageDraw.Draw(text_mask)
    for word in words:
        try:
            left = float(word["left"]) * scan_w
            right = float(word["right"]) * scan_w
            top = float(word["top"]) * scan_h
            bottom = float(word["bottom"]) * scan_h
        except (KeyError, TypeError, ValueError):
            continue
        if right <= left or bottom <= top:
            continue
        pad_x = max(2, (right - left) * 0.25)
        pad_y = max(2, (bottom - top) * 0.45)
        text_draw.rectangle(
            [
                max(0, left - pad_x),
                max(0, top - pad_y),
                min(scan_w, right + pad_x),
                min(scan_h, bottom + pad_y),
            ],
            fill=255,
        )
    text_px = text_mask.load()

    x_min_scan = int(scan_w * 0.04)
    x_max_scan = int(scan_w * 0.96)
    row_ranges: list[tuple[float, float, int, float, float]] = []
    active: list[bool] = []
    for y in range(scan_h):
        xs = []
        for x in range(x_min_scan, x_max_scan):
            if text_px[x, y]:
                continue
            if gray_px[x, y] < 235:
                xs.append(x)
        if not xs:
            row_ranges.append((0.0, 0.0, 0, 0.0, 0.0))
            active.append(False)
            continue
        count = len(xs)
        left = min(xs) / scan_w
        right = max(xs) / scan_w
        density = count / max(scan_w, 1)
        span = right - left
        row_ranges.append((left, right, count, density, span))
        active.append(
            (density > 0.012 and span > 0.28)
            or density > 0.045
            or (span > 0.55 and count > 6)
        )

    bridge = max(6, int(scan_h * 0.02))
    i = 0
    while i < scan_h:
        if active[i]:
            i += 1
            continue
        j = i
        while j < scan_h and not active[j]:
            j += 1
        if i > 0 and j < scan_h and (j - i) <= bridge:
            for k in range(i, j):
                active[k] = True
        i = j

    broad_bands: list[tuple[float, float]] = []
    y = 0
    while y < scan_h:
        if not active[y]:
            y += 1
            continue
        y0 = y
        while y < scan_h and active[y]:
            y += 1
        y1 = y
        if (y1 - y0) / scan_h >= 0.035:
            broad_bands.append((y0 / scan_h, y1 / scan_h))

    if not broad_bands:
        return []

    lines = _ocr_word_lines(words)
    heights = sorted(max(float(line["bottom"]) - float(line["top"]), 0.001) for line in lines)
    median_h = heights[len(heights) // 2] if heights else 0.012
    gap_threshold = max(0.035, median_h * 3.0)

    clusters: list[list[dict]] = []
    cur: list[dict] = []
    for line in lines:
        if not cur:
            cur = [line]
            continue
        gap = float(line["top"]) - float(cur[-1]["bottom"])
        if gap > gap_threshold:
            clusters.append(cur)
            cur = [line]
        else:
            cur.append(line)
    if cur:
        clusters.append(cur)

    rects: list[fitz.Rect] = []

    def foreground_x_range(y0_frac: float, y1_frac: float) -> tuple[float, float] | None:
        y0 = max(0, int(y0_frac * scan_h))
        y1 = min(scan_h, int(y1_frac * scan_h))
        xs: list[float] = []
        for yy in range(y0, y1):
            left, right, count, _, span = row_ranges[yy]
            if count and span > 0.05:
                xs.extend([left, right])
        if not xs:
            return None
        return max(0.04, min(xs) - 0.01), min(0.96, max(xs) + 0.01)

    if len(words) <= 8:
        for band_top, band_bottom in broad_bands:
            xr = foreground_x_range(band_top, band_bottom)
            if not xr:
                continue
            rect = fitz.Rect(xr[0] * page_w, band_top * page_h, xr[1] * page_w, band_bottom * page_h)
            if rect.width >= page_w * 0.25 and rect.height >= page_h * 0.18:
                rects.append(_expanded_clip(rect, page_rect, margin=4))
        return rects

    for cluster_idx, cluster in enumerate(clusters):
        if cluster_idx + 1 >= len(clusters):
            continue
        c_top = float(cluster[0]["top"])
        c_bottom = float(cluster[-1]["bottom"])
        prev_bottom = float(clusters[cluster_idx - 1][-1]["bottom"]) if cluster_idx > 0 else 0.045
        next_top = float(clusters[cluster_idx + 1][0]["top"]) if cluster_idx + 1 < len(clusters) else 0.95
        gap_after = next_top - c_bottom
        gap_before = c_top - prev_bottom
        if gap_after < gap_threshold * 1.25 and gap_before < gap_threshold * 1.25:
            continue

        visual_top = max(0.045, prev_bottom + 0.006 if gap_before >= gap_threshold else c_top - median_h * 2.0)
        visual_bottom = min(0.95, next_top - 0.006 if gap_after >= gap_threshold else c_bottom + median_h * 2.0)
        visual_top = max(0.045, visual_top - max(0.012, median_h * 1.5))
        if visual_bottom - visual_top < 0.08:
            continue

        overlaps_band = any(not (visual_bottom < b0 or visual_top > b1) for b0, b1 in broad_bands)
        if not overlaps_band:
            continue
        xr = foreground_x_range(visual_top, visual_bottom)
        if not xr:
            continue
        rect = fitz.Rect(xr[0] * page_w, visual_top * page_h, xr[1] * page_w, visual_bottom * page_h)
        if rect.width < page_w * 0.18 or rect.height < page_h * 0.18:
            continue
        area_frac = rect.get_area() / max(page_rect.get_area(), 1)
        if area_frac < 0.20 or area_frac > 0.78:
            continue
        if _count_body_text_lines_in_rect(rect, page, words) >= 4:
            continue
        if any(_rect_overlap_ratio(rect, existing) > 0.55 for existing in rects):
            continue
        rects.append(_expanded_clip(rect, page_rect, margin=4))

    return sorted(rects, key=lambda r: (r.y0, r.x0))


def _scan_page_visual_envelope(
    page: fitz.Page,
    words: list[dict],
    render_matrix: fitz.Matrix,
) -> list[fitz.Rect]:
    """Find the outer boundary of large sparse scanned drawings.

    Density and connected-component scans can crop architectural drawings to
    their darkest sub-region. This pass masks OCR text and estimates one large
    envelope from the remaining foreground pixels.
    """
    page_rect = page.rect
    page_w = max(page_rect.width, 1)
    page_h = max(page_rect.height, 1)
    try:
        pix = page.get_pixmap(matrix=render_matrix, alpha=False)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    except Exception:
        return []

    scale = min(1.0, 900 / max(img.width, img.height))
    scan_w = max(1, int(img.width * scale))
    scan_h = max(1, int(img.height * scale))
    scan = img.resize((scan_w, scan_h), Image.Resampling.BILINEAR) if scale < 1 else img

    text_mask = Image.new("L", (scan_w, scan_h), 0)
    text_draw = ImageDraw.Draw(text_mask)
    for word in words:
        try:
            left = float(word["left"]) * scan_w
            right = float(word["right"]) * scan_w
            top = float(word["top"]) * scan_h
            bottom = float(word["bottom"]) * scan_h
        except (KeyError, TypeError, ValueError):
            continue
        if right <= left or bottom <= top:
            continue
        pad_x = max(2, (right - left) * 0.18)
        pad_y = max(2, (bottom - top) * 0.32)
        text_draw.rectangle(
            [max(0, left - pad_x), max(0, top - pad_y), min(scan_w, right + pad_x), min(scan_h, bottom + pad_y)],
            fill=255,
        )

    gray = scan.convert("L")
    rgb = scan.load()
    gray_px = gray.load()
    text_px = text_mask.load()
    xs: list[int] = []
    ys: list[int] = []
    col_bins: set[int] = set()
    row_bins: set[int] = set()
    for y in range(scan_h):
        for x in range(scan_w):
            if text_px[x, y]:
                continue
            r, g, b = rgb[x, y]
            lum = gray_px[x, y]
            sat = max(r, g, b) - min(r, g, b)
            if lum < 224 or sat > 24:
                xs.append(x)
                ys.append(y)
                col_bins.add(min(7, x * 8 // max(scan_w, 1)))
                row_bins.add(min(7, y * 8 // max(scan_h, 1)))

    fg_count = len(xs)
    total_px = max(scan_w * scan_h, 1)
    fg_frac = fg_count / total_px
    if fg_frac < 0.0012 or fg_frac > 0.38:
        return []
    if len(col_bins) < 4 or len(row_bins) < 4:
        return []

    xs.sort()
    ys.sort()

    def pick(values: list[int], frac: float) -> int:
        idx = int(round((len(values) - 1) * frac))
        return values[max(0, min(len(values) - 1, idx))]

    x0 = pick(xs, 0.004)
    x1 = pick(xs, 0.996)
    y0 = pick(ys, 0.004)
    y1 = pick(ys, 0.996)
    if x1 <= x0 or y1 <= y0:
        return []

    rect = fitz.Rect(
        x0 / scan_w * page_w,
        y0 / scan_h * page_h,
        x1 / scan_w * page_w,
        y1 / scan_h * page_h,
    )
    area_frac = rect.get_area() / max(page_rect.get_area(), 1)
    if area_frac < 0.16 or area_frac > 0.78:
        return []
    if rect.width < page_w * 0.45 or rect.height < page_h * 0.30:
        return []
    if rect.y0 < page_h * 0.035 and rect.height < page_h * 0.18:
        return []
    if rect.y0 > page_h * 0.82 and rect.height < page_h * 0.16:
        return []
    if _count_body_text_lines_in_rect(rect, page, words) >= 4:
        return []

    return [_expanded_clip(rect, page_rect, margin=max(4, min(page_w, page_h) * 0.012))]


def _remove_scan_overlapped_infos(
    infos: list[dict],
    saved_rects: list[fitz.Rect],
    rect: fitz.Rect,
    job_id: str,
) -> None:
    remove_indices = [
        i for i, existing in enumerate(saved_rects)
        if _rect_overlap_min_ratio(rect, existing) > 0.58 and rect.get_area() > existing.get_area() * 1.10
    ]
    for i in reversed(remove_indices):
        image_id = str(infos[i].get("image_id", ""))
        if image_id:
            try:
                _job_image_path(job_id, image_id).unlink(missing_ok=True)
            except Exception:
                pass
        del infos[i]
        del saved_rects[i]


def _refine_scan_gap_rect(
    page: fitz.Page,
    words: list[dict],
    fig_top: float,
    fig_bot: float,
    render_matrix: fitz.Matrix,
) -> fitz.Rect | None:
    """Refine coarse OCR gap bands to a tighter visual bbox using rendered pixels."""
    page_rect = page.rect
    page_w = max(page_rect.width, 1)
    page_h = max(page_rect.height, 1)

    try:
        pix = page.get_pixmap(matrix=render_matrix, alpha=False)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    except Exception:
        return None

    if img.width <= 0 or img.height <= 0:
        return None

    scale = min(1.0, 960 / max(img.width, img.height))
    scan_w = max(1, int(img.width * scale))
    scan_h = max(1, int(img.height * scale))
    scan = img.resize((scan_w, scan_h), Image.Resampling.BILINEAR) if scale < 1 else img

    gy0 = max(0, min(scan_h - 1, int(fig_top * scan_h)))
    gy1 = max(gy0 + 1, min(scan_h, int(fig_bot * scan_h)))
    if gy1 - gy0 < max(10, int(scan_h * 0.03)):
        return None

    text_mask = Image.new("L", (scan_w, scan_h), 0)
    text_draw = ImageDraw.Draw(text_mask)
    for word in words:
        try:
            left = float(word["left"]) * scan_w
            right = float(word["right"]) * scan_w
            top = float(word["top"]) * scan_h
            bottom = float(word["bottom"]) * scan_h
        except (KeyError, TypeError, ValueError):
            continue
        if right <= left or bottom <= top:
            continue
        pad_x = max(2, (right - left) * 0.16)
        pad_y = max(2, (bottom - top) * 0.33)
        text_draw.rectangle(
            [
                max(0, left - pad_x),
                max(0, top - pad_y),
                min(scan_w, right + pad_x),
                min(scan_h, bottom + pad_y),
            ],
            fill=255,
        )

    gray = scan.convert("L")
    rgb = scan.load()
    gray_px = gray.load()
    text_px = text_mask.load()

    def _is_foreground(x: int, y: int) -> bool:
        if text_px[x, y]:
            return False
        r, g, b = rgb[x, y]
        lum = gray_px[x, y]
        sat = max(r, g, b) - min(r, g, b)
        return lum < 232 or sat > 24

    row_density: list[float] = []
    for y in range(gy0, gy1):
        hits = 0
        for x in range(scan_w):
            if _is_foreground(x, y):
                hits += 1
        row_density.append(hits / max(scan_w, 1))

    active_rows = [i for i, density in enumerate(row_density) if density > 0.010]
    if not active_rows:
        return None
    ry0 = gy0 + active_rows[0]
    ry1 = gy0 + active_rows[-1] + 1
    if ry1 - ry0 < max(8, int(scan_h * 0.02)):
        return None

    col_density: list[float] = []
    row_span = max(1, ry1 - ry0)
    for x in range(scan_w):
        hits = 0
        for y in range(ry0, ry1):
            if _is_foreground(x, y):
                hits += 1
        col_density.append(hits / row_span)

    active_cols = [i for i, density in enumerate(col_density) if density > 0.020]
    if not active_cols:
        return None
    cx0 = active_cols[0]
    cx1 = active_cols[-1] + 1
    if cx1 - cx0 < max(10, int(scan_w * 0.10)):
        return None

    margin_x = max(3, int(scan_w * 0.006))
    margin_y = max(3, int(scan_h * 0.006))
    rect = fitz.Rect(
        max(0, cx0 - margin_x) / scan_w * page_w,
        max(0, ry0 - margin_y) / scan_h * page_h,
        min(scan_w, cx1 + margin_x) / scan_w * page_w,
        min(scan_h, ry1 + margin_y) / scan_h * page_h,
    )
    area_frac = rect.get_area() / max(page_rect.get_area(), 1.0)
    if rect.width < page_w * 0.12 or rect.height < page_h * 0.06:
        return None
    if area_frac > 0.88:
        return None
    return rect


def _extract_ocr_figure_images(
    pdf_path: Path,
    page_words_list: list[list[dict]],
    job_id: str,
) -> list[list[dict]]:
    """For scanned PDFs: detect figure regions from large vertical gaps in OCR text layout.

    Renders any region where OCR found no text for >= 15% of the page height.
    Ignores pure-margin gaps (at very top/bottom of page with no adjacent text).
    """
    if not job_id or not page_words_list:
        return [[] for _ in page_words_list]
    try:
        doc = fitz.open(str(pdf_path))
    except Exception:
        return [[] for _ in page_words_list]

    result: list[list[dict]] = []
    mat = fitz.Matrix(2, 2)
    MIN_GAP = 0.15      # 15% of page height minimum
    TOP_SKIP = 0.05     # ignore gaps entirely above this y (header margin)
    BOT_SKIP = 0.95     # ignore gaps entirely below this y (footer margin)

    for page_idx, words in enumerate(page_words_list):
        if page_idx >= doc.page_count:
            result.append([])
            continue

        page = doc[page_idx]
        page_rect = page.rect
        page_h = max(page_rect.height, 1)
        page_w = max(page_rect.width, 1)

        _word_mids = [(w["left"] + w["right"]) / 2 for w in words
                      if w.get("right", 0) > w.get("left", 0)]
        _lw = sum(1 for x in _word_mids if x < 0.46)
        _rw = sum(1 for x in _word_mids if x > 0.54)
        _two_col_ocr = (
            len(_word_mids) >= 10
            and _lw >= 5 and _rw >= 5
            and min(_lw, _rw) / len(_word_mids) > 0.20
        ) if _word_mids else False

        def _ocr_order_key(r: fitz.Rect) -> float:
            yf = max(0.0, min(1.0, r.y0 / page_h))
            if not _two_col_ocr or (r.x1 - r.x0) > page_w * 0.55:
                return yf
            col = 0 if (r.x0 + r.x1) / 2 < page_w * 0.5 else 1
            return (col + yf) / 2

        infos: list[dict] = []

        if not words:
            result.append(infos)
            continue

        # Sort all word boxes by their top edge, add sentinels at 0 and 1
        sorted_words = sorted(words, key=lambda w: w["top"])
        sentinels = (
            [{"top": 0.0, "bottom": 0.0}]
            + sorted_words
            + [{"top": 1.0, "bottom": 1.0}]
        )

        used: list[tuple[float, float]] = []
        saved_rects: list[fitz.Rect] = []
        gap_candidate_rects: list[fitz.Rect] = []
        for i in range(len(sentinels) - 1):
            gap_top = sentinels[i]["bottom"]
            gap_bot = sentinels[i + 1]["top"]
            gap_size = gap_bot - gap_top
            if gap_size < MIN_GAP:
                continue
            # Skip pure-header or pure-footer gaps
            if gap_bot <= TOP_SKIP or gap_top >= BOT_SKIP:
                continue
            # Clip gap to visible area
            fig_top = max(gap_top, TOP_SKIP)
            fig_bot = min(gap_bot, BOT_SKIP)
            if fig_bot - fig_top < MIN_GAP * 0.8:
                continue
            # Require meaningful text context to distinguish figures from blank pages.
            words_above = sum(1 for w in words if w["bottom"] <= gap_top + 0.01)
            words_below = sum(1 for w in words if w["top"] >= gap_bot - 0.01)
            if words_above < 3:
                total_context_words = words_above + words_below
                if total_context_words <= 5:
                    pass  # near-photo-only page ??let brightness/mid-gray check decide
                elif gap_top > 0.05 or words_below < 10:
                    # Allow top-of-page photos: gap starts at very top with lots of text below.
                    continue
            # For mid-page gaps also require text below (otherwise it's trailing whitespace).
            # Gaps that extend to the bottom 80%+ of the page are allowed with text only above.
            if words_below < 3 and fig_bot < 0.82:
                continue
            # Deduplicate overlapping regions
            if any(not (fig_bot <= u0 or fig_top >= u1) for u0, u1 in used):
                continue
            used.append((fig_top, fig_bot))

            refined = _refine_scan_gap_rect(page, words, fig_top, fig_bot, mat)
            if refined is None:
                continue
            if any(
                _rect_overlap_min_ratio(refined, existing) > 0.62
                for existing in gap_candidate_rects
            ):
                continue
            clip = _expanded_clip(
                refined,
                page_rect,
                margin=2,
            )
            if _count_body_text_lines_in_rect(clip, page, words) >= 3:
                continue
            try:
                pix = page.get_pixmap(matrix=mat, clip=clip, alpha=False)
            except Exception:
                continue
            # Skip blank or text-heavy bands; keep actual visuals and line drawings.
            img_check = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            sample = list(img_check.convert("L").getdata())
            if not sample:
                continue
            total_px = len(sample)
            mean_brightness = sum(sample) / total_px
            mid_gray_count = sum(1 for p in sample if 100 <= p <= 200)
            dark_count = sum(1 for p in sample if p <= 90)
            mid_gray_frac = mid_gray_count / total_px
            dark_frac = dark_count / total_px
            if mean_brightness > 238:   # near-white blank region
                continue
            if mid_gray_frac < 0.08 and dark_frac < 0.03:
                continue
            gap_candidate_rects.append(fitz.Rect(clip))

        for rect in _ocr_caption_visual_rects(page, words):
            _remove_scan_overlapped_infos(infos, saved_rects, rect, job_id)
            if any(_rect_overlap_ratio(rect, existing) > 0.35 or _rect_has_meaningful_overlap(rect, existing) for existing in saved_rects):
                continue
            clip = _expanded_clip(rect, page_rect, margin=3)
            try:
                pix = page.get_pixmap(matrix=mat, clip=clip, alpha=False)
            except Exception:
                continue
            image_id = f"scan_cap_p{page_idx}_c{len(infos)}.png"
            img_path = _job_image_path(job_id, image_id)
            img_path.parent.mkdir(parents=True, exist_ok=True)
            pix.save(str(img_path))
            saved_rects.append(fitz.Rect(clip))
            y_frac = max(0.0, min(1.0, rect.y0 / page_h))
            infos.append({
                "image_id": image_id,
                "y_frac": y_frac,
                "order_key": _ocr_order_key(clip),
                "left": clip.x0 / page_w,
                "right": clip.x1 / page_w,
                "top": clip.y0 / page_h,
                "bottom": clip.y1 / page_h,
            })

        for rect in _scan_page_visual_envelope(page, words, mat):
            _remove_scan_overlapped_infos(infos, saved_rects, rect, job_id)
            if any(_rect_overlap_ratio(rect, existing) > 0.35 or _rect_has_meaningful_overlap(rect, existing) for existing in saved_rects):
                continue
            clip = _expanded_clip(rect, page_rect, margin=3)
            try:
                pix = page.get_pixmap(matrix=mat, clip=clip, alpha=False)
            except Exception:
                continue
            image_id = f"scan_page_p{page_idx}_e{len(infos)}.png"
            img_path = _job_image_path(job_id, image_id)
            img_path.parent.mkdir(parents=True, exist_ok=True)
            pix.save(str(img_path))
            saved_rects.append(fitz.Rect(clip))
            y_frac = max(0.0, min(1.0, rect.y0 / page_h))
            infos.append({
                "image_id": image_id,
                "y_frac": y_frac,
                "order_key": _ocr_order_key(clip),
                "left": clip.x0 / page_w,
                "right": clip.x1 / page_w,
                "top": clip.y0 / page_h,
                "bottom": clip.y1 / page_h,
            })

        for rect in _scan_layout_visual_rects(page, words, mat):
            _remove_scan_overlapped_infos(infos, saved_rects, rect, job_id)
            if any(_rect_overlap_ratio(rect, existing) > 0.35 or _rect_has_meaningful_overlap(rect, existing) for existing in saved_rects):
                continue
            clip = _expanded_clip(rect, page_rect, margin=3)
            try:
                pix = page.get_pixmap(matrix=mat, clip=clip, alpha=False)
            except Exception:
                continue
            image_id = f"scan_layout_p{page_idx}_l{len(infos)}.png"
            img_path = _job_image_path(job_id, image_id)
            img_path.parent.mkdir(parents=True, exist_ok=True)
            pix.save(str(img_path))
            saved_rects.append(fitz.Rect(clip))
            y_frac = max(0.0, min(1.0, rect.y0 / page_h))
            infos.append({
                "image_id": image_id,
                "y_frac": y_frac,
                "order_key": _ocr_order_key(clip),
                "left": clip.x0 / page_w,
                "right": clip.x1 / page_w,
                "top": clip.y0 / page_h,
                "bottom": clip.y1 / page_h,
            })

        for rect in _scan_density_visual_rects(page, words, mat):
            _remove_scan_overlapped_infos(infos, saved_rects, rect, job_id)
            if any(_rect_overlap_ratio(rect, existing) > 0.35 for existing in saved_rects):
                continue
            clip = _expanded_clip(rect, page_rect, margin=3)
            try:
                pix = page.get_pixmap(matrix=mat, clip=clip, alpha=False)
            except Exception:
                continue
            image_id = f"scan_den_p{page_idx}_d{len(infos)}.png"
            img_path = _job_image_path(job_id, image_id)
            img_path.parent.mkdir(parents=True, exist_ok=True)
            pix.save(str(img_path))
            saved_rects.append(fitz.Rect(clip))
            y_frac = max(0.0, min(1.0, rect.y0 / page_h))
            infos.append({
                "image_id": image_id,
                "y_frac": y_frac,
                "order_key": _ocr_order_key(clip),
                "left": clip.x0 / page_w,
                "right": clip.x1 / page_w,
                "top": clip.y0 / page_h,
                "bottom": clip.y1 / page_h,
            })

        for rect in _scan_visual_rects_from_render(page, words, mat):
            _remove_scan_overlapped_infos(infos, saved_rects, rect, job_id)
            if any(_rect_overlap_ratio(rect, existing) > 0.35 or _rect_has_meaningful_overlap(rect, existing) for existing in saved_rects):
                continue
            clip = _expanded_clip(rect, page_rect, margin=3)
            try:
                pix = page.get_pixmap(matrix=mat, clip=clip, alpha=False)
            except Exception:
                continue
            image_id = f"scan_img_p{page_idx}_r{len(infos)}.png"
            img_path = _job_image_path(job_id, image_id)
            img_path.parent.mkdir(parents=True, exist_ok=True)
            pix.save(str(img_path))
            saved_rects.append(fitz.Rect(clip))
            y_frac = max(0.0, min(1.0, rect.y0 / page_h))
            infos.append({
                "image_id": image_id,
                "y_frac": y_frac,
                "order_key": _ocr_order_key(clip),
                "left": clip.x0 / page_w,
                "right": clip.x1 / page_w,
                "top": clip.y0 / page_h,
                "bottom": clip.y1 / page_h,
            })

        # Lowest-priority fallback: OCR vertical-gap candidates.
        # Keep only those that do not conflict with tighter render/caption detections.
        for rect in gap_candidate_rects:
            if any(_rect_overlap_ratio(rect, existing) > 0.35 or _rect_has_meaningful_overlap(rect, existing) for existing in saved_rects):
                continue
            clip = _expanded_clip(rect, page_rect, margin=1)
            try:
                pix = page.get_pixmap(matrix=mat, clip=clip, alpha=False)
            except Exception:
                continue
            image_id = f"scan_gap_p{page_idx}_g{len(infos)}.png"
            img_path = _job_image_path(job_id, image_id)
            img_path.parent.mkdir(parents=True, exist_ok=True)
            pix.save(str(img_path))
            saved_rects.append(fitz.Rect(clip))
            y_frac = max(0.0, min(1.0, rect.y0 / page_h))
            infos.append({
                "image_id": image_id,
                "y_frac": y_frac,
                "order_key": _ocr_order_key(clip),
                "left": clip.x0 / page_w,
                "right": clip.x1 / page_w,
                "top": clip.y0 / page_h,
                "bottom": clip.y1 / page_h,
            })

        infos.sort(key=lambda item: float(item.get("order_key", item.get("y_frac", 0))))
        result.append(infos)

    doc.close()
    return result


def _call_gemini(client: genai.Client, model: str, prompt: str, retries: int = 4, job_id: str | None = None) -> str:
    import concurrent.futures
    for attempt in range(retries):
        try:
            with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(
                    client.models.generate_content, model=model, contents=prompt
                )
                try:
                    response = future.result(timeout=180)
                except concurrent.futures.TimeoutError:
                    raise TimeoutError("Gemini API 請求超時，請稍後再試")
            return response.text.strip()
        except Exception as e:
            err = str(e)
            is_retryable = any(code in err for code in ["429", "503", "UNAVAILABLE", "RESOURCE_EXHAUSTED", "timed out", "timeout", "TimeoutError", "DeadlineExceeded"])
            if is_retryable and attempt < retries - 1:
                wait = 20 * (attempt + 1)
                print(f"[{model}] retryable error; waiting {wait} seconds before retry")
                if job_id:
                    _set_job(job_id, message=f"Gemini is busy; retrying in {wait} seconds")
                time.sleep(wait)
            else:
                if is_retryable:
                    _cooling_models[model] = time.time() + 600  # cooldown 10 min
                    print(f"[{model}] retryable error; cooling down for 10 minutes")
                raise


def _validate_model(client: genai.Client, model: str) -> bool:
    """Send a minimal request to check if model is usable. Returns False only on 404."""
    if model in _dead_models:
        return False
    try:
        _call_gemini(client, model, "hi", retries=1)
        return True
    except Exception as e:
        if "404" in str(e) or "NOT_FOUND" in str(e):
            _dead_models.add(model)
            print(f"[{model}] model returned 404; removing from candidates")
            return False
        return True  # 503/429 = temporarily busy, still usable


def _is_chinese_target(target_lang: str) -> bool:
    target = (target_lang or "").lower()
    return any(token in target for token in ("chinese", "traditional", "simplified", "中文", "繁體", "简体"))


def _text_script_counts(text: str) -> tuple[int, int]:
    latin = len(re.findall(r"[A-Za-z]", text or ""))
    cjk = sum(1 for ch in (text or "") if 0x4E00 <= ord(ch) <= 0x9FFF)
    return latin, cjk


def _looks_untranslated_output(source_text: str, translated_text: str, target_lang: str) -> bool:
    if not _is_chinese_target(target_lang):
        return False
    src_latin, _ = _text_script_counts(source_text)
    out_latin, out_cjk = _text_script_counts(translated_text)
    if src_latin < 80 or out_latin < 50:
        return False
    # Proper nouns and citations can remain Latin, but a Chinese translation should
    # still contain substantial CJK text for prose-heavy chunks.
    if out_cjk < max(20, int(out_latin * 0.35)):
        return True

    consecutive_english_lines = 0
    for line in _normalize_text(translated_text).split("\n"):
        clean = line.strip()
        if not clean:
            consecutive_english_lines = 0
            continue
        latin, cjk = _text_script_counts(clean)
        english_heavy = latin >= 18 and cjk < max(3, int(latin * 0.15))
        if english_heavy:
            consecutive_english_lines += 1
            if consecutive_english_lines >= 3:
                return True
        else:
            consecutive_english_lines = 0
    return False


def _text_similarity_key(text: str) -> str:
    clean = _normalize_text(text or "").lower()
    return "".join(
        ch for ch in clean
        if ch.isalnum() or 0x4E00 <= ord(ch) <= 0x9FFF
    )


def _texts_are_near_duplicates(a: str, b: str, min_len: int = 80, threshold: float = 0.92) -> bool:
    ka = _text_similarity_key(a)
    kb = _text_similarity_key(b)
    if min(len(ka), len(kb)) < min_len:
        return False
    if ka == kb:
        return True
    return SequenceMatcher(None, ka, kb).ratio() >= threshold


def _native_translation_needs_retry(source_text: str, translated_text: str, target_lang: str) -> bool:
    """Catch native-PDF batch marker drift before it reaches overlay/export."""
    translated_text = _normalize_text(translated_text or "").strip()
    if not translated_text:
        return True
    if _looks_untranslated_output(source_text, translated_text, target_lang):
        return True
    if not _is_chinese_target(target_lang):
        return False
    source_text = _normalize_text(source_text or "").strip()
    src_latin, _ = _text_script_counts(source_text)
    _, out_cjk = _text_script_counts(translated_text)
    # A short native block producing a much longer Chinese paragraph usually means
    # the marker batch merged in content from the next block.
    return src_latin >= 90 and len(source_text) <= 300 and out_cjk > max(110, int(len(source_text) * 0.85))


def _translate_chunk(client: genai.Client, text: str, target_lang: str, model: str | None = None, job_id: str | None = None) -> str:
    prompt = f"""You are a professional translator and document specialist.

Task: Translate the following document text into {target_lang}.

Rules:
- Translate ALL content into {target_lang} regardless of the source language ??English, Japanese, Chinese, Korean, or any other language must ALL be translated into {target_lang}. Never leave any source language text untranslated.
- Preserve ALL paragraph breaks and line spacing
- Keep titles, headings, and section structure intact
- Maintain bullet points, numbering, and indentation
- Translate accurately and naturally ??avoid word-for-word literal translation
- Keep proper nouns, technical terms, and abbreviations appropriate for the target language
- Output ONLY the translated text, no explanations or notes

Text to translate:
{text}"""

    primary = model or GEMINI_MODEL
    fallbacks = sorted([m for m in GEMINI_FALLBACK_MODELS if m != primary], key=_model_cost_rank)
    models_to_try = [primary] + fallbacks
    last_err = None
    for m in models_to_try:
        if m in _dead_models or _is_cooling(m):
            continue
        try:
            translated = _call_gemini(client, m, prompt, job_id=job_id)
            if _looks_untranslated_output(text, translated, target_lang):
                retry_prompt = f"""The previous translation output left too much English untranslated.

Translate the text below into {target_lang} again.

Hard requirements:
- Translate every English prose sentence into {target_lang}.
- Do not copy English sentences unchanged.
- Keep paragraph and line breaks aligned with the input as much as possible.
- Keep formulas, equation numbers, units, names, journal abbreviations, and citations only where they should remain literal.
- Output only the translated document text.

Text:
{text}"""
                if job_id:
                    _set_job(job_id, message="Gemini 回傳內容不完整，正在重試翻譯...")
                translated_retry = _call_gemini(client, m, retry_prompt, retries=2, job_id=job_id)
                if not _looks_untranslated_output(text, translated_retry, target_lang):
                    return _normalize_text(translated_retry)
            return _normalize_text(translated)
        except Exception as e:
            err = str(e)
            if "404" in err or "NOT_FOUND" in err:
                _dead_models.add(m)
                print(f"[{m}] model returned 404; trying next model")
            else:
                print(f"[{m}] translation failed; trying next model ({e})")
            last_err = e
    raise last_err


def _parse_scanned_translation(text: str, expected_count: int) -> dict[int, str]:
    """Parse [Pn] marker output into {para_idx: translated_text}."""
    import re
    result: dict[int, str] = {}
    matches = list(re.finditer(r"\[P(\d+)\]", text, re.IGNORECASE))
    if not matches:
        return {}
    for i, match in enumerate(matches):
        idx = int(match.group(1))
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        para_text = text[start:end].strip()
        result[idx] = para_text
    return result


def _translate_chunk_scanned(
    client: genai.Client,
    para_texts: list[str],
    target_lang: str,
    model: str | None = None,
    job_id: str | None = None,
) -> dict[int, str]:
    """Translate paragraphs with [Pn] position markers preserved in output."""
    if not para_texts:
        return {}

    marked_parts = []
    for i, t in enumerate(para_texts):
        marked_parts.append(f"[P{i}]\n{t.strip()}")
    text = "\n\n".join(marked_parts)

    prompt = f"""You are a professional translator and document specialist.

Task: Translate the following document paragraphs into {target_lang}.

CRITICAL RULE: Each paragraph begins with a position marker like [P0], [P1], [P2], etc.
You MUST keep these markers EXACTLY as-is at the START of each translated paragraph.
Do NOT remove, alter, merge, or add any [Pn] markers.

Rules:
- Translate ALL content into {target_lang}
- Keep [Pn] marker at the start of the corresponding translated paragraph
- Translate accurately and naturally
- Keep proper nouns, citations, and technical terms as appropriate
- Output ONLY the translated text with [Pn] markers, no explanations

Text to translate:
{text}"""

    primary = model or GEMINI_MODEL
    fallbacks = sorted([m for m in GEMINI_FALLBACK_MODELS if m != primary], key=_model_cost_rank)
    models_to_try = [primary] + fallbacks
    last_err = None
    for m in models_to_try:
        if m in _dead_models or _is_cooling(m):
            continue
        try:
            raw = _call_gemini(client, m, prompt, job_id=job_id)
            parsed = _parse_scanned_translation(raw, len(para_texts))
            if len(parsed) >= max(1, len(para_texts) // 2):
                return parsed
            # Retry if too few markers preserved
            retry_prompt = (
                prompt
                + "\n\nRe-translate and ensure every paragraph starts with its [Pn] marker. "
                "Output must begin with [P0]."
            )
            raw2 = _call_gemini(client, m, retry_prompt, retries=2, job_id=job_id)
            parsed2 = _parse_scanned_translation(raw2, len(para_texts))
            return parsed2 if len(parsed2) >= len(parsed) else parsed
        except Exception as e:
            err = str(e)
            if "404" in err or "NOT_FOUND" in err:
                _dead_models.add(m)
                print(f"[{m}] model returned 404; trying next model")
            else:
                print(f"[{m}] scanned translation failed; trying next model ({e})")
            last_err = e
    if last_err:
        raise last_err
    return {}


def _translate_pages_scanned(
    client: genai.Client,
    pages: list[str],
    page_words: list[list[dict]],
    page_lines: list[list[dict]],
    target_lang: str,
    job_id: str,
    model: str | None = None,
    on_page_done=None,
    page_fig_blocks: list[list[dict]] | None = None,
) -> tuple[list[str], list[list[dict]], list[list[dict]]]:
    """Like _translate_pages but produces para_idx-aware translated_blocks.

    Returns (translated_pages, translated_blocks, page_ocr_paras).
    Falls back to chunk translation per page if paragraph grouping fails.
    page_fig_blocks: per-page figure bboxes (normalized); words inside are excluded.
    """
    translated_pages: list[str] = []
    translated_blocks_all: list[list[dict]] = []
    page_ocr_paras: list[list[dict]] = []
    total = len(pages)

    def _outside_figs(items: list[dict], figs: list[dict]) -> list[dict]:
        if not figs:
            return items
        result = []
        for item in items:
            cy = (item["top"] + item["bottom"]) / 2
            cx = (item["left"] + item["right"]) / 2
            in_fig = any(
                fig["top"] <= cy <= fig["bottom"] and fig["left"] <= cx <= fig["right"]
                for fig in figs
                if "top" in fig and "bottom" in fig and "left" in fig and "right" in fig
            )
            if not in_fig:
                result.append(item)
        return result

    def _ocr_words_to_text(words: list[dict]) -> str:
        lines = _ocr_word_lines(words)
        return _normalize_text("\n".join(
            str(line.get("text", "")).strip()
            for line in lines
            if str(line.get("text", "")).strip()
        ))

    for i, page_text in enumerate(pages, 1):
        _set_job(job_id, message=f"Translating page {i}/{total}", progress=50 + int(40 * i / total))

        plines = page_lines[i - 1] if i - 1 < len(page_lines) else []
        pwords = page_words[i - 1] if i - 1 < len(page_words) else []
        pfigs = page_fig_blocks[i - 1] if page_fig_blocks and i - 1 < len(page_fig_blocks) else []
        plines = _outside_figs(plines, pfigs)
        pwords = _outside_figs(pwords, pfigs)
        filtered_page_text = _ocr_words_to_text(pwords)

        if not filtered_page_text.strip():
            translated_pages.append("")
            translated_blocks_all.append([])
            page_ocr_paras.append([])
            if on_page_done:
                on_page_done(translated_pages, translated_blocks_all, page_ocr_paras)
            if i < total:
                time.sleep(PAGE_DELAY_SECONDS)
            continue

        paras = _group_ocr_paragraphs(plines)

        if not paras:
            # No para grouping — fall back to chunk translation
            chunks = [filtered_page_text[j: j + 3000] for j in range(0, len(filtered_page_text), 3000)]
            translated_chunks = [_translate_chunk(client, c, target_lang, model=model, job_id=job_id) for c in chunks]
            trans = _normalize_text("\n".join(translated_chunks))
            translated_pages.append(trans)
            translated_blocks_all.append(_text_to_translation_blocks(trans))
            page_ocr_paras.append([])
        else:
            para_texts = _assign_words_to_paragraphs(pwords, paras)
            nonempty_para_indices = [pi for pi, text in enumerate(para_texts) if text.strip()]
            if not nonempty_para_indices:
                translated_pages.append("")
                translated_blocks_all.append([])
                page_ocr_paras.append(paras)
                if on_page_done:
                    on_page_done(translated_pages, translated_blocks_all, page_ocr_paras)
                if i < total:
                    time.sleep(PAGE_DELAY_SECONDS)
                continue

            try:
                compact_para_texts = [para_texts[pi] for pi in nonempty_para_indices]
                compact_para_map = _translate_chunk_scanned(client, compact_para_texts, target_lang, model=model, job_id=job_id)
                para_map = {
                    nonempty_para_indices[compact_pi]: text
                    for compact_pi, text in compact_para_map.items()
                    if 0 <= compact_pi < len(nonempty_para_indices)
                }
            except Exception as exc:
                print(f"[scanned translation] page {i} failed ({exc}), falling back to filtered OCR chunk")
                chunks = [filtered_page_text[j: j + 3000] for j in range(0, len(filtered_page_text), 3000)]
                translated_chunks = [_translate_chunk(client, c, target_lang, model=model, job_id=job_id) for c in chunks]
                trans = _normalize_text("\n".join(translated_chunks))
                translated_pages.append(trans)
                # Set para_idx even in fallback path to help frontend overlay mapping
                fallback_blocks = _text_to_translation_blocks(trans)
                for block_index, b in enumerate(fallback_blocks):
                    if b.get("type") == "text":
                        b["para_idx"] = block_index
                translated_blocks_all.append(fallback_blocks)
                page_ocr_paras.append(paras)
                if on_page_done:
                    on_page_done(translated_pages, translated_blocks_all, page_ocr_paras)
                if i < total:
                    time.sleep(PAGE_DELAY_SECONDS)
                continue

            blocks: list[dict] = []
            parts: list[str] = []
            for pi in range(len(paras)):
                ptext = _normalize_text(para_map.get(pi, para_texts[pi] or ""))
                blocks.append({"type": "text", "para_idx": pi, "source_idx": round(pi * TRANSLATION_BLOCK_STEP, 6), "text": ptext})
                parts.append(ptext)

            translated_pages.append("\n\n".join(parts))
            translated_blocks_all.append(blocks)
            page_ocr_paras.append(paras)

        if on_page_done:
            on_page_done(translated_pages, translated_blocks_all, page_ocr_paras)
        if i < total:
            time.sleep(PAGE_DELAY_SECONDS)

    return translated_pages, translated_blocks_all, page_ocr_paras


def _translate_pages_native_blocks(
    client: genai.Client,
    pages: list[str],
    page_source_blocks: list[list[dict]],
    target_lang: str,
    job_id: str,
    model: str | None = None,
    on_page_done=None,
) -> tuple[list[str], list[list[dict]]]:
    """Translate native PDF text block-by-block so overlay positions stay aligned."""
    translated_pages: list[str] = []
    translated_blocks_all: list[list[dict]] = []
    total = len(pages)

    for i, page_text in enumerate(pages, 1):
        _set_job(job_id, message=f"Translating page {i}/{total}", progress=50 + int(40 * i / total))

        source_blocks = page_source_blocks[i - 1] if i - 1 < len(page_source_blocks) else []
        text_blocks = [
            block for block in source_blocks
            if isinstance(block, dict) and str(block.get("text", "")).strip()
        ]

        if not text_blocks:
            if not page_text.strip():
                translated_pages.append("")
                translated_blocks_all.append([])
            else:
                chunks = [page_text[j: j + 3000] for j in range(0, len(page_text), 3000)]
                translated_chunks = [_translate_chunk(client, chunk, target_lang, model=model, job_id=job_id) for chunk in chunks]
                trans = _normalize_text("\n".join(translated_chunks))
                translated_pages.append(trans)
                translated_blocks_all.append(_text_to_translation_blocks(trans))
            if on_page_done:
                on_page_done(translated_pages, translated_blocks_all, [])
            if i < total:
                time.sleep(PAGE_DELAY_SECONDS)
            continue

        para_texts = [_normalize_text(str(block.get("text", ""))).strip() for block in text_blocks]
        try:
            para_map = _translate_chunk_scanned(client, para_texts, target_lang, model=model, job_id=job_id)
            if len(para_map) < max(1, int(len(para_texts) * 0.75)):
                raise ValueError("too few preserved native paragraph markers")
        except Exception as exc:
            print(f"[native translation] page {i} marker translation failed ({exc}), falling back to page chunk")
            chunks = [page_text[j: j + 3000] for j in range(0, len(page_text), 3000)]
            translated_chunks = [_translate_chunk(client, chunk, target_lang, model=model, job_id=job_id) for chunk in chunks]
            trans = _normalize_text("\n".join(translated_chunks))
            split = _split_translation_text_for_paras(trans, len(para_texts))
            para_map = {pi: split[pi] for pi in range(min(len(split), len(para_texts)))}

        blocks: list[dict] = []
        parts: list[str] = []
        resolved_texts: list[str] = []
        for pi, source_block in enumerate(text_blocks):
            ptext = _normalize_text(str(para_map.get(pi, ""))).strip()
            duplicate_of_previous = any(
                _texts_are_near_duplicates(ptext, prev_text)
                and not _texts_are_near_duplicates(para_texts[pi], para_texts[prev_i], min_len=120, threshold=0.80)
                for prev_i, prev_text in enumerate(resolved_texts)
            )
            if (
                _native_translation_needs_retry(para_texts[pi], ptext, target_lang)
                or duplicate_of_previous
            ):
                try:
                    ptext = _translate_chunk(
                        client,
                        para_texts[pi],
                        target_lang,
                        model=model,
                        job_id=job_id,
                    ).strip()
                except Exception as exc:
                    print(f"[native translation] page {i} block {pi} retry failed ({exc})")
                    if not ptext:
                        ptext = para_texts[pi]
            resolved_texts.append(ptext)
            try:
                source_idx = float(source_block.get("source_idx", pi * TRANSLATION_BLOCK_STEP))
            except Exception:
                source_idx = pi * TRANSLATION_BLOCK_STEP
            blocks.append({
                "type": "text",
                "para_idx": pi,
                "source_idx": round(source_idx, 6),
                "text": ptext,
            })
            parts.append(ptext)

        translated_pages.append("\n\n".join(parts))
        translated_blocks_all.append(blocks)

        if on_page_done:
            on_page_done(translated_pages, translated_blocks_all, [])
        if i < total:
            time.sleep(PAGE_DELAY_SECONDS)

    return translated_pages, translated_blocks_all


def _translate_pages(client: genai.Client, pages: list[str], target_lang: str, job_id: str, model: str | None = None, on_page_done=None) -> list[str]:
    translated = []
    total = len(pages)
    for i, page_text in enumerate(pages, 1):
        _set_job(job_id, message=f"Translating page {i}/{total}", progress=50 + int(40 * i / total))
        if not page_text.strip():
            translated.append("")
        else:
            chunks = [page_text[j: j + 3000] for j in range(0, len(page_text), 3000)]
            translated_chunks = [_translate_chunk(client, chunk, target_lang, model=model, job_id=job_id) for chunk in chunks]
            translated.append(_normalize_text("\n".join(translated_chunks)))
        if on_page_done:
            on_page_done(translated)
        if i < total:
            time.sleep(PAGE_DELAY_SECONDS)
    return translated


def _render_translation_blocks_to_cell(job_id: str | None, cell, blocks: list[dict]):
    paragraphs = list(cell.paragraphs)
    for paragraph in paragraphs[1:]:
        p = paragraph._element
        p.getparent().remove(p)

    if not paragraphs:
        paragraphs = [cell.add_paragraph()]

    first_para = paragraphs[0]
    first_para.clear()

    if not blocks:
        first_para.add_run("(No translation)")
        return

    for index, block in enumerate(blocks):
        paragraph = first_para if index == 0 else cell.add_paragraph()
        paragraph.clear()

        if block.get("type") == "image":
            if job_id is None:
                continue
            image_id = str(block.get("image_id", "")).strip()
            image_path = _job_image_path(job_id, image_id)
            if not image_id or not image_path.exists():
                paragraph.add_run("[Image]")
                continue
            run = paragraph.add_run()
            user_w = int(block.get("width", 0) or 0)
            img_width = Inches(min(user_w / 96, 5.2)) if user_w > 0 else Inches(5.2)
            run.add_picture(str(image_path), width=img_width)
            continue

        paragraph.add_run(_normalize_text(str(block.get("text", ""))) or "")


def _build_compare_word(
    original_pages: list[str],
    translated_pages: list[str],
    output_path: Path,
    translated_blocks: list[list[dict]] | None = None,
    job_id: str | None = None,
    font_size: int = 10,
):
    doc = Document()
    title = doc.add_heading("PDF OCR Translation", level=1)
    title.alignment = 1

    total_pages = max(len(original_pages), len(translated_pages))
    for i in range(total_pages):
        orig = original_pages[i] if i < len(original_pages) else ""
        trans = translated_pages[i] if i < len(translated_pages) else ""
        doc.add_heading(f"Page {i + 1}", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Original"
        hdr_cells[1].text = "Translation"

        for cell in hdr_cells:
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell.paragraphs[0].paragraph_format.alignment = 1
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "2E74B5")
            tcPr.append(shd)

        row_cells = table.add_row().cells
        row_cells[0].text = _normalize_text(orig) if orig else "(No original text)"
        page_blocks = translated_blocks[i] if translated_blocks and i < len(translated_blocks) else None
        if page_blocks is None:
            row_cells[1].text = _normalize_text(trans) if trans else "(No translation)"
        else:
            _render_translation_blocks_to_cell(job_id, row_cells[1], page_blocks)

        for cell in row_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(font_size)

        doc.add_paragraph()

    doc.save(str(output_path))


def _build_reading_word(
    translated_pages: list[str],
    output_path: Path,
    translated_blocks: list[list[dict]] | None = None,
    job_id: str | None = None,
    font_size: int = 10,
):
    doc = Document()
    _apply_reading_doc_style(doc, font_size)

    flat_blocks = list(_iter_translation_blocks(translated_pages, translated_blocks))
    first_text = ""
    next_text = ""
    for block in flat_blocks:
        if block.get("type") == "text" and _normalize_text(str(block.get("text", ""))).strip():
            if not first_text:
                first_text = _normalize_text(str(block.get("text", ""))).strip()
            else:
                next_text = _normalize_text(str(block.get("text", ""))).strip()
                break

    should_promote_title = _looks_like_title(first_text, next_text)
    title_consumed = False

    for block in flat_blocks:
        block_type = block.get("type", "text")
        if block_type == "image":
            if job_id is None:
                continue
            image_id = str(block.get("image_id", "")).strip()
            image_path = _job_image_path(job_id, image_id)
            if not image_id or not image_path.exists():
                continue

            image_para = doc.add_paragraph()
            image_para.alignment = 1
            image_para.paragraph_format.space_before = Pt(6)
            image_para.paragraph_format.space_after = Pt(3)
            image_run = image_para.add_run()
            user_w = int(block.get("width", 0) or 0)
            img_width = Inches(min(user_w / 96, 5.5)) if user_w > 0 else Inches(5.5)
            image_run.add_picture(str(image_path), width=img_width)

            caption = _normalize_text(str(block.get("name", ""))).strip()
            if caption:
                caption_para = doc.add_paragraph()
                caption_para.alignment = 1
                caption_para.paragraph_format.space_after = Pt(10)
                caption_run = caption_para.add_run(caption)
                _style_run(caption_run, size=Pt(max(8, font_size - 2)))
            continue

        text = _normalize_text(str(block.get("text", "")))
        if should_promote_title and not title_consumed and text.strip() == first_text:
            title_para = doc.add_paragraph()
            title_para.paragraph_format.space_after = Pt(font_size + 4)
            title_run = title_para.add_run(first_text)
            _style_run(title_run, size=Pt(font_size + 4), bold=True)
            title_consumed = True
            continue

        if not text.strip():
            doc.add_paragraph()
            continue

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
        run = paragraph.add_run(text)
        blk_size = max(8, min(72, int(block.get("font_size") or font_size)))
        blk_bold = bool(block.get("bold", False))
        _style_run(run, size=Pt(blk_size), bold=blk_bold)

    doc.save(str(output_path))


def _build_original_layout_word(pdf_path: Path, output_path: Path, original_pages: list[str] | None = None):
    if not pdf_path.exists():
        raise FileNotFoundError(f"Original PDF not found: {pdf_path}")

    doc = Document()
    with fitz.open(str(pdf_path)) as pdf_doc:
        if pdf_doc.page_count == 0:
            doc.save(str(output_path))
            return

        first_rect = pdf_doc[0].rect
        section = doc.sections[0]
        section.page_width = Inches(first_rect.width / 72)
        section.page_height = Inches(first_rect.height / 72)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        image_width = section.page_width - section.left_margin - section.right_margin

        for index, page in enumerate(pdf_doc):
            if index > 0:
                doc.add_page_break()

            # Page A: original PDF rendered as image
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image_stream = BytesIO(pixmap.tobytes("png"))
            img_para = doc.add_paragraph()
            img_para.alignment = 1
            img_para.paragraph_format.space_before = Pt(0)
            img_para.paragraph_format.space_after = Pt(0)
            max_height = section.page_height - section.top_margin - section.bottom_margin
            aspect = pixmap.width / pixmap.height
            img_run = img_para.add_run()
            if image_width / aspect > max_height:
                img_run.add_picture(image_stream, height=max_height)
            else:
                img_run.add_picture(image_stream, width=image_width)

            # Page B: OCR text of this page
            heading = doc.add_paragraph()
            heading.paragraph_format.page_break_before = True
            heading.paragraph_format.space_after = Pt(8)
            heading_run = heading.add_run(f"Page {index + 1} OCR text")
            _style_run(heading_run, size=Pt(13), bold=True)

            ocr_text = (original_pages[index] if original_pages and index < len(original_pages) else "").strip()
            if ocr_text:
                for line in _normalize_text(ocr_text).split("\n"):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(4)
                    _style_run(p.add_run(line), size=Pt(10))
            else:
                p = doc.add_paragraph()
                _style_run(p.add_run("(No OCR text)"), size=Pt(10))

    doc.save(str(output_path))


def _page_to_image_bytes(page: fitz.Page) -> tuple[bytes, str]:
    """Return (image_bytes, ext) for use as the page thumbnail in reading-compare.

    For scanned PDFs the page is just a large embedded raster — extract it at its
    native resolution instead of re-rendering.  For native vector PDFs fall back to
    rendering at 3× (≈216 DPI) which is sharper than the previous 2×.
    """
    page_area = page.rect.get_area()
    images = page.get_images(full=True)
    best: tuple[int, bytes, str] | None = None
    for img_info in images:
        xref = img_info[0]
        try:
            d = page.parent.extract_image(xref)
        except Exception:
            continue
        if not d or not d.get("image"):
            continue
        w, h = d.get("width", 0), d.get("height", 0)
        if w * h < page_area * 0.5:  # skip decorative small images
            continue
        if best is None or w * h > best[0]:
            best = (w * h, d["image"], d.get("ext", "png"))
    if best:
        return best[1], best[2]
    pix = page.get_pixmap(matrix=fitz.Matrix(3, 3), alpha=False)
    return pix.tobytes("png"), "png"


def _build_reading_compare_word(
    pdf_path: Path,
    output_path: Path,
    translated_pages: list[str],
    translated_blocks: list[list[dict]] | None = None,
    job_id: str | None = None,
    font_size: int = 10,
):
    if not pdf_path.exists():
        raise FileNotFoundError(f"Original PDF not found: {pdf_path}")

    doc = Document()
    _apply_reading_doc_style(doc, font_size)

    with fitz.open(str(pdf_path)) as pdf_doc:
        if pdf_doc.page_count == 0:
            doc.save(str(output_path))
            return

        first_rect = pdf_doc[0].rect
        section = doc.sections[0]
        section.page_width = Inches(first_rect.width / 72)
        section.page_height = Inches(first_rect.height / 72)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        content_w_emu, content_h_emu = _section_content_size_emu(section)
        translation_img_max_w_emu = min(int(Inches(5.5)), content_w_emu)

        total = pdf_doc.page_count
        for index, page in enumerate(pdf_doc):
            if index > 0:
                doc.add_page_break()

            # Page A: original PDF as image
            img_bytes, _img_ext = _page_to_image_bytes(page)
            image_stream = BytesIO(img_bytes)
            img_para = doc.add_paragraph()
            img_para.alignment = 1
            img_para.paragraph_format.space_before = Pt(0)
            img_para.paragraph_format.space_after = Pt(0)
            img_run = img_para.add_run()
            _add_picture_within_bounds(
                img_run,
                image_stream,
                content_w_emu,
                content_h_emu,
                safety_pt=12.0,
            )

            # Page B: translated reading text
            doc.add_page_break()
            page_blocks = translated_blocks[index] if translated_blocks and index < len(translated_blocks) else None
            page_text = translated_pages[index] if index < len(translated_pages) else ""

            if page_blocks:
                for block in page_blocks:
                    if block.get("type") == "image":
                        if job_id:
                            image_id = str(block.get("image_id", "")).strip()
                            image_path = _job_image_path(job_id, image_id)
                            if image_id and image_path.exists():
                                p = doc.add_paragraph()
                                p.alignment = 1
                                p.paragraph_format.space_before = Pt(4)
                                p.paragraph_format.space_after = Pt(6)
                                _add_picture_within_bounds(
                                    p.add_run(),
                                    str(image_path),
                                    translation_img_max_w_emu,
                                    content_h_emu,
                                    safety_pt=10.0,
                                )
                        continue
                    text = _normalize_text(str(block.get("text", "")))
                    if not text.strip():
                        doc.add_paragraph()
                        continue
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(8)
                    blk_size = max(8, min(72, int(block.get("font_size") or font_size)))
                    blk_bold = bool(block.get("bold", False))
                    _style_run(p.add_run(text), size=Pt(blk_size), bold=blk_bold)
            elif page_text.strip():
                for line in _normalize_text(page_text).split("\n"):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(6)
                    _style_run(p.add_run(line), size=Pt(font_size))
            else:
                p = doc.add_paragraph()
                _style_run(p.add_run("(No translation)"), size=Pt(font_size))

    doc.save(str(output_path))


def _build_word(
    original_pages: list[str],
    translated_pages: list[str],
    output_path: Path,
    translated_blocks: list[list[dict]] | None = None,
    job_id: str | None = None,
    mode: str = "compare",
    pdf_path: Path | None = None,
    font_size: int = 10,
):
    if mode == "reading":
        _build_reading_word(
            translated_pages,
            output_path,
            translated_blocks=translated_blocks,
            job_id=job_id,
            font_size=font_size,
        )
        return

    if mode == "original_layout":
        if pdf_path is None:
            raise FileNotFoundError("Original PDF path is required for original layout download")
        _build_original_layout_word(pdf_path, output_path, original_pages=original_pages)
        return

    if mode == "reading_compare":
        if pdf_path is None:
            raise FileNotFoundError("Original PDF path is required for reading compare download")
        _build_reading_compare_word(
            pdf_path, output_path,
            translated_pages=translated_pages,
            translated_blocks=translated_blocks,
            job_id=job_id,
            font_size=font_size,
        )
        return

    _build_compare_word(
        original_pages,
        translated_pages,
        output_path,
        translated_blocks=translated_blocks,
        job_id=job_id,
        font_size=font_size,
    )



# ---------------------------------------------------------------------------
# Background worker
# ---------------------------------------------------------------------------

def _process_job(
    job_id: str,
    pdf_path: Path,
    ocr_lang: str,
    target_lang: str,
    api_key: str,
    source_name: str,
    gemini_model: str | None = None,
):
    try:
        _set_job(job_id, status="running", progress=5, message="Reading PDF")

        client = _get_client(api_key)
        active_model = gemini_model or GEMINI_MODEL
        _set_job(job_id, progress=7, message=f"Checking model {active_model}")
        if not _validate_model(client, active_model):
            _set_job(job_id, status="error",
                     message=f"Model {active_model} is unavailable. Please choose another model.",
                     dead_model=active_model)
            return
        gemini_model = active_model

        _set_job(job_id, progress=10, message="Extracting PDF text")
        # First pass without job_id so we don't save scan backgrounds as figure images.
        # If OCR is not needed we re-run with job_id to extract real inline figures.
        native_pages, native_page_lines, native_page_words, native_source_blocks, _native_table_blocks_no_img = _extract_text_native(pdf_path, job_id=None)

        is_scanned = _needs_ocr(native_pages)
        if is_scanned:
            _set_job(job_id, progress=20, message="Running OCR")
            pages, page_lines, page_words = _extract_text_ocr(pdf_path, lang=ocr_lang)
            page_source_blocks: list[list[dict]] = []
            # Detect figure regions from OCR word positions (large text-free gaps)
            ocr_fig_blocks = _extract_ocr_figure_images(pdf_path, page_words, job_id)
            # For scanned PDFs, native_table_blocks contains page-background scan images,
            # not content figures.  Use only OCR gap-detected regions.
            page_table_blocks = ocr_fig_blocks
        else:
            _set_job(job_id, progress=20, message="Using embedded PDF text")
            # Re-run native extraction with job_id:
            # 1) save inline figure assets to disk
            # 2) exclude figure / formula / table regions from extracted text used for translation
            native_pages_with_visuals, native_page_lines_with_visuals, native_page_words_with_visuals, native_source_blocks_with_visuals, native_table_blocks = _extract_text_native(pdf_path, job_id=job_id)
            pages = native_pages_with_visuals
            page_lines = native_page_lines_with_visuals
            page_words = native_page_words_with_visuals
            page_source_blocks = native_source_blocks_with_visuals
            page_table_blocks = native_table_blocks

        total_chars = sum(len(p) for p in pages)
        _set_job(job_id, progress=40,
                 message=f"Extracted {len(pages)} pages, {total_chars} characters. Starting translation.")

        def _on_page_done(
            translated_so_far: list[str],
            blocks_so_far: list[list[dict]] | None = None,
            ocr_paras_so_far: list[list[dict]] | None = None,
        ) -> None:
            partial_blocks = _clone_translation_blocks(blocks_so_far) if blocks_so_far is not None else _pages_to_translation_blocks(translated_so_far)
            if not is_scanned and page_table_blocks:
                partial_blocks = _merge_table_blocks(_clone_translation_blocks(partial_blocks), page_table_blocks)
            partial_ocr_paras = [[dict(p) for p in page] for page in ocr_paras_so_far] if ocr_paras_so_far is not None else []
            partial = {
                "original": pages,
                "translated": translated_so_far,
                "page_lines": page_lines,
                "page_words": page_words,
                "page_source_blocks": page_source_blocks,
                "page_ocr_paras": partial_ocr_paras,
                "translated_blocks": partial_blocks,
                "translated_blueprints": _clone_translation_blocks(partial_blocks),
                "page_table_blocks": page_table_blocks,
                "target_lang": target_lang,
                "source_name": source_name,
                "gemini_model": gemini_model or GEMINI_MODEL,
                "updated_at": time.time(),
                "partial": True,
            }
            with STORE_LOCK:
                job_results[job_id] = partial
            _save_partial_result(job_id, partial)

        if is_scanned:
            translated_pages, translated_blocks, page_ocr_paras = _translate_pages_scanned(
                client, pages, page_words, page_lines, target_lang, job_id,
                model=gemini_model, on_page_done=_on_page_done,
                page_fig_blocks=page_table_blocks,
            )
            translated_blocks = _merge_table_blocks(translated_blocks, page_table_blocks)
        else:
            page_ocr_paras: list[list[dict]] = []
            translated_pages, translated_blocks = _translate_pages_native_blocks(
                client,
                pages,
                page_source_blocks,
                target_lang,
                job_id,
                model=gemini_model,
                on_page_done=_on_page_done,
            )
            translated_blocks = _merge_table_blocks(translated_blocks, page_table_blocks)

        _set_job(job_id, progress=92, message="Building Word output")
        output_path = OUTPUT_FOLDER / f"{job_id}.docx"
        _build_word(
            pages,
            translated_pages,
            output_path,
            translated_blocks=translated_blocks,
            job_id=job_id,
        )

        saved_pdf = OUTPUT_FOLDER / f"{job_id}.pdf"
        pdf_path.rename(saved_pdf)

        final_data = {
            "original": pages,
            "translated": translated_pages,
            "page_lines": page_lines,
            "page_words": page_words,
            "page_source_blocks": page_source_blocks,
            "page_ocr_paras": page_ocr_paras,
            "translated_blocks": translated_blocks,
            "translated_blueprints": _clone_translation_blocks(translated_blocks),
            "page_table_blocks": page_table_blocks,
            "target_lang": target_lang,
            "source_name": source_name,
            "updated_at": time.time(),
        }
        with STORE_LOCK:
            job_results[job_id] = final_data
        _save_partial_result(job_id, final_data)

        _set_job(job_id, status="done", progress=100, message="Done",
                 download_url=f"/download/{job_id}")
    except Exception as exc:
        _set_job(job_id, status="error", message=str(exc))
        try:
            pdf_path.unlink(missing_ok=True)
        except Exception:
            pass


# ---------------------------------------------------------------------------
def _resume_job(job_id: str, api_key: str):
    persist_repaired = None
    with STORE_LOCK:
        data = job_results.get(job_id)
        if data and _repair_scanned_overlay_data(data):
            persist_repaired = deepcopy(data)
    if not data:
        return
    if persist_repaired is not None:
        _save_partial_result(job_id, persist_repaired)

    original = data.get("original", [])
    translated_so_far = list(data.get("translated", []))
    page_lines = data.get("page_lines", [])
    page_words = data.get("page_words", [])
    page_source_blocks = data.get("page_source_blocks", [])
    page_ocr_paras = data.get("page_ocr_paras", [])
    translated_blocks_so_far = data.get("translated_blocks", _pages_to_translation_blocks(translated_so_far))
    page_table_blocks = data.get("page_table_blocks", [])
    target_lang = data.get("target_lang", TARGET_LANGUAGE)
    source_name = data.get("source_name", "output")
    gemini_model = data.get("gemini_model") or GEMINI_MODEL
    is_scanned_resume = (
        isinstance(page_ocr_paras, list)
        and any(isinstance(page, list) and page for page in page_ocr_paras)
        and not any(isinstance(page, list) and page for page in page_source_blocks)
    )

    remaining = original[len(translated_so_far):]
    if not remaining:
        _set_job(job_id, status="done", progress=100, message="Done",
                 download_url=f"/download/{job_id}")
        return

    try:
        _set_job(job_id, status="running", progress=50, message=f"Resuming translation from page {len(translated_so_far)+1}")
        client = _get_client(api_key)
        total = len(original)
        existing_blocks = _clone_translation_blocks(translated_blocks_so_far[:len(translated_so_far)])
        while len(existing_blocks) < len(translated_so_far):
            existing_blocks.append(_text_to_translation_blocks(translated_so_far[len(existing_blocks)]))
        existing_ocr_paras = [[dict(p) for p in page] for page in page_ocr_paras[:len(translated_so_far)]]

        def _save_resume_partial(
            combined: list[str],
            combined_blocks: list[list[dict]],
            combined_ocr_paras: list[list[dict]] | None = None,
        ) -> None:
            blocks_copy = _clone_translation_blocks(combined_blocks)
            partial = {
                "original": original,
                "translated": combined,
                "page_lines": page_lines,
                "page_words": page_words,
                "page_source_blocks": page_source_blocks,
                "page_ocr_paras": [[dict(p) for p in page] for page in (combined_ocr_paras or [])],
                "translated_blocks": blocks_copy,
                "translated_blueprints": _clone_translation_blocks(blocks_copy),
                "page_table_blocks": page_table_blocks,
                "target_lang": target_lang,
                "source_name": source_name,
                "gemini_model": gemini_model,
                "updated_at": time.time(),
                "partial": True,
            }
            with STORE_LOCK:
                job_results[job_id] = partial
            _save_partial_result(job_id, partial)

        def _on_page_done(translated_all: list[str]) -> None:
            combined = translated_so_far + translated_all
            _save_resume_partial(
                combined,
                existing_blocks + _pages_to_translation_blocks(translated_all),
                [],
            )

        def _set_job_resume(translated_all: list[str]) -> None:
            done = len(translated_so_far) + len(translated_all)
            _set_job(job_id, message=f"Translating page {done}/{total}",
                     progress=50 + int(40 * done / total))
            _on_page_done(translated_all)

        if is_scanned_resume:
            start = len(translated_so_far)

            def _set_job_resume_scanned(
                translated_all: list[str],
                blocks_all: list[list[dict]],
                ocr_paras_all: list[list[dict]],
            ) -> None:
                done = start + len(translated_all)
                _set_job(job_id, message=f"Translating page {done}/{total}",
                         progress=50 + int(40 * done / total))
                _save_resume_partial(
                    translated_so_far + translated_all,
                    existing_blocks + _clone_translation_blocks(blocks_all),
                    existing_ocr_paras + [[dict(p) for p in page] for page in ocr_paras_all],
                )

            new_translated, new_blocks, new_ocr_paras = _translate_pages_scanned(
                client,
                remaining,
                page_words[start:],
                page_lines[start:],
                target_lang,
                job_id,
                model=gemini_model,
                on_page_done=_set_job_resume_scanned,
                page_fig_blocks=page_table_blocks[start:] if isinstance(page_table_blocks, list) else None,
            )
            all_translated = translated_so_far + new_translated
            all_blocks = existing_blocks + _clone_translation_blocks(new_blocks)
            all_ocr_paras = existing_ocr_paras + [[dict(p) for p in page] for page in new_ocr_paras]
            all_blocks = _merge_table_blocks(all_blocks, page_table_blocks)
        else:
            new_translated = _translate_pages(
                client, remaining, target_lang, job_id,
                model=gemini_model, on_page_done=_set_job_resume,
            )
            all_translated = translated_so_far + new_translated
            all_blocks = _merge_table_blocks(existing_blocks + _pages_to_translation_blocks(new_translated), page_table_blocks)
            all_ocr_paras = []

        _set_job(job_id, progress=92, message="Building Word output")
        output_path = OUTPUT_FOLDER / f"{job_id}.docx"
        _build_word(original, all_translated, output_path, translated_blocks=all_blocks, job_id=job_id)

        final_data = {
            "original": original,
            "translated": all_translated,
            "page_lines": page_lines,
            "page_words": page_words,
            "page_source_blocks": page_source_blocks,
            "page_ocr_paras": all_ocr_paras,
            "translated_blocks": all_blocks,
            "translated_blueprints": _clone_translation_blocks(all_blocks),
            "page_table_blocks": page_table_blocks,
            "target_lang": target_lang,
            "source_name": source_name,
            "updated_at": time.time(),
        }
        with STORE_LOCK:
            job_results[job_id] = final_data
        _save_partial_result(job_id, final_data)
        _set_job(job_id, status="done", progress=100, message="Done",
                 download_url=f"/download/{job_id}")
    except Exception as exc:
        _set_job(job_id, status="error", message=str(exc))


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.after_request
def _no_store_dynamic_routes(response):
    path = request.path or ""
    if path == "/" or path == "/jobs" or path.startswith(("/result/", "/status/", "/jobs/")):
        response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
    return response


@app.route("/")
def index():
    response = app.make_response(render_template("index.html", target_language=TARGET_LANGUAGE))
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response


@app.route("/cover/<filename>")
def cover(filename: str):
    cover_dir = _BUNDLE / "cover"
    return send_from_directory(str(cover_dir), filename)


@app.route("/favicon.ico")
def favicon():
    cover_dir = _BUNDLE / "cover"
    return send_from_directory(str(cover_dir), "raw toas..ico", mimetype="image/x-icon")


@app.route("/resume/<job_id>", methods=["POST"])
def resume(job_id: str):
    body = request.get_json(force=True) or {}
    api_key = (body.get("api_key") or GEMINI_API_KEY or "").strip()
    if not api_key:
        return jsonify({"error": "請輸入 Gemini API Key"}), 400
    with STORE_LOCK:
        data = job_results.get(job_id)
    if not data:
        return jsonify({"error": "找不到此翻譯工作"}), 404
    if not data.get("partial"):
        return jsonify({"error": "此翻譯工作已完成，無需繼續翻譯"}), 400
    _set_job(job_id, status="running", progress=50, message="Preparing to resume translation")
    thread = threading.Thread(target=_resume_job, args=(job_id, api_key), daemon=True)
    thread.start()
    return jsonify({"job_id": job_id})


@app.route("/upload", methods=["GET"])
def upload_page_redirect():
    return redirect("/")


@app.route("/upload", methods=["POST"])
def upload():
    if "file" not in request.files:
        return jsonify({"error": "Missing file"}), 400

    file = request.files["file"]
    if not file.filename or not file.filename.lower().endswith(".pdf"):
        return jsonify({"error": "請上傳 PDF 檔案"}), 400

    api_key = request.form.get("api_key", "").strip() or GEMINI_API_KEY
    if not api_key:
        return jsonify({"error": "請輸入 Gemini API Key"}), 400

    ocr_lang = request.form.get("ocr_lang", "chi_tra+chi_sim+eng")
    target_lang = request.form.get("target_lang", TARGET_LANGUAGE)
    gemini_model = request.form.get("gemini_model", "").strip() or None
    source_name = Path(file.filename).stem

    job_id = str(uuid.uuid4())
    pdf_path = UPLOAD_FOLDER / f"{job_id}.pdf"
    file.save(str(pdf_path))

    _set_job(job_id, status="queued", progress=0, message="Queued", source_name=source_name)

    thread = threading.Thread(
        target=_process_job,
        args=(job_id, pdf_path, ocr_lang, target_lang, api_key, source_name),
        kwargs={"gemini_model": gemini_model},
        daemon=True,
    )
    thread.start()

    print(f"[upload] accepted {source_name}.pdf as job {job_id}")
    if request.headers.get("X-Requested-With") != "fetch":
        return (
            "<!doctype html><html><head><meta charset='utf-8'>"
            "<title>ButterLayer</title></head><body>"
            "<script>"
            f"localStorage.setItem('butterlayer.activeJobId', {json.dumps(job_id)});"
            "location.replace('/');"
            "</script>"
            "<p>Returning to ButterLayer...</p>"
            "</body></html>"
        )
    return jsonify({"job_id": job_id, "source_name": source_name})


@app.route("/status/<job_id>")
def status(job_id: str):
    job = _get_job(job_id)
    if job is None:
        return jsonify({"error": "找不到此翻譯工作"}), 404
    return jsonify(job)


@app.route("/fetch-models", methods=["POST"])
def fetch_models():
    import urllib.request, urllib.error
    body = request.get_json(force=True) or {}
    api_key = body.get("api_key", "").strip()
    if not api_key:
        return jsonify({"error": "請輸入 API Key"}), 400
    try:
        url = f"https://generativelanguage.googleapis.com/v1beta/models?key={api_key}&pageSize=100"
        req = urllib.request.Request(url, headers={"Accept": "application/json"})
        with urllib.request.urlopen(req, timeout=15) as resp:
            data = json.loads(resp.read().decode())
        _exclude = ("image", "tts", "audio", "vision", "embedding", "aqa", "bison", "gecko", "robotics", "computer-use")
        models = []
        for m in data.get("models", []):
            methods = m.get("supportedGenerationMethods", [])
            name = m.get("name", "")
            if "generateContent" not in methods or "gemini" not in name:
                continue
            model_id = name.replace("models/", "")
            if any(kw in model_id for kw in _exclude):
                continue
            if model_id in _dead_models:
                continue
            models.append({"id": model_id})
        models.sort(key=lambda x: x["id"])
        return jsonify({"models": models})
    except urllib.error.HTTPError as e:
        return jsonify({"error": f"API request failed ({e.code}). Please check whether the API Key is valid."}), 400
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/jobs")
def list_jobs():
    result_list = []
    try:
        for path in OUTPUT_FOLDER.glob("*_result.json"):
            job_id = path.stem.replace("_result", "")
            try:
                mtime = path.stat().st_mtime
                cached = job_results.get(job_id)
                if cached:
                    source_name = cached.get("source_name", job_id)
                    partial = bool(cached.get("partial"))
                else:
                    with open(str(path), "r", encoding="utf-8") as f:
                        d = json.load(f)
                    source_name = d.get("source_name", job_id)
                    partial = bool(d.get("partial"))
                result_list.append({"id": job_id, "name": source_name, "date": mtime, "partial": partial})
            except Exception:
                continue
    except Exception:
        pass
    result_list.sort(key=lambda x: x["date"], reverse=True)
    return jsonify(result_list)


@app.route("/jobs/<job_id>", methods=["DELETE"])
def delete_job(job_id: str):
    import shutil
    if not all(c in "abcdefghijklmnopqrstuvwxyz0123456789-_" for c in job_id.lower()):
        return jsonify({"error": "Invalid job_id"}), 400
    with STORE_LOCK:
        job_results.pop(job_id, None)
        jobs.pop(job_id, None)
    for suffix in ["_result.json", ".pdf", ".docx",
                   "_reading.docx", "_reading_compare.docx",
                   "_original_layout.docx", "_compare.docx"]:
        p = OUTPUT_FOLDER / f"{job_id}{suffix}"
        if p.exists():
            try: p.unlink()
            except Exception: pass
    assets = OUTPUT_FOLDER / f"{job_id}_assets"
    if assets.exists():
        shutil.rmtree(str(assets), ignore_errors=True)
    return jsonify({"ok": True})


@app.route("/result/<job_id>")
def result(job_id: str):
    try:
        data = _get_result(job_id)
    except Exception as exc:
        return jsonify({"error": f"載入結果失敗：{exc}"}), 500
    if data is None:
        return jsonify({"error": "找不到此翻譯記錄，可能已被刪除或伺服器已重啟。"}), 404
    return jsonify(data)


@app.route("/result/<job_id>/source_blocks")
def result_source_blocks(job_id: str):
    """Re-extract page_source_blocks from the saved PDF (backfills old jobs)."""
    with STORE_LOCK:
        data = job_results.get(job_id)
    if data is None:
        return jsonify({"error": "Result not found"}), 404
    # If already present, return immediately
    if data.get("page_source_blocks"):
        return jsonify({"page_source_blocks": data["page_source_blocks"]})
    pdf_path = OUTPUT_FOLDER / f"{job_id}.pdf"
    if not pdf_path.exists():
        return jsonify({"page_source_blocks": []}), 200
    try:
        doc = fitz.open(str(pdf_path))
        blocks = []
        for page in doc:
            blocks.append(_extract_source_block_positions(page))
        doc.close()
        with STORE_LOCK:
            job_results[job_id]["page_source_blocks"] = blocks
        _save_partial_result(job_id, job_results[job_id])
        return jsonify({"page_source_blocks": blocks})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/result/<job_id>/translation", methods=["POST"])
def save_translation(job_id: str):
    body = request.get_json(force=True) or {}
    page = body.get("page")
    try:
        page = int(page)
    except (TypeError, ValueError):
        return jsonify({"error": "Invalid page number"}), 400

    if page < 1:
        return jsonify({"error": "Page number must be greater than 0"}), 400

    text = _normalize_text(body.get("text", ""))
    blocks = _normalize_translation_blocks(body.get("blocks"), text, job_id=job_id)
    text = _translation_blocks_to_text(blocks)

    persist_data = None
    with STORE_LOCK:
        data = job_results.get(job_id)
        if data is None:
            return jsonify({"error": "Result not found"}), 404

        translated = data.get("translated")
        translated_blocks = data.setdefault(
            "translated_blocks",
            _pages_to_translation_blocks(data.get("translated", [])),
        )
        data.setdefault(
            "translated_blueprints",
            _clone_translation_blocks(translated_blocks),
        )
        if not isinstance(translated, list) or page > len(translated):
            return jsonify({"error": "Page number is out of range"}), 400
        if not isinstance(translated_blocks, list) or page > len(translated_blocks):
            return jsonify({"error": "Page number is out of range"}), 400

        translated[page - 1] = text
        translated_blocks[page - 1] = blocks
        updated_at = time.time()
        data["updated_at"] = updated_at
        persist_data = deepcopy(data)

    if persist_data is not None:
        _save_partial_result(job_id, persist_data)

    return jsonify({"ok": True, "page": page, "updated_at": updated_at})


@app.route("/result/<job_id>/image", methods=["POST"])
def upload_translation_image(job_id: str):
    with STORE_LOCK:
        if job_id not in job_results:
            return jsonify({"error": "Result not found"}), 404

    if "image" not in request.files:
        return jsonify({"error": "Missing image"}), 400

    image_file = request.files["image"]
    if not image_file.filename:
        return jsonify({"error": "Missing image"}), 400

    ext = Path(image_file.filename).suffix.lower()
    if ext not in ALLOWED_IMAGE_EXTENSIONS:
        return jsonify({"error": "Unsupported image type"}), 400

    image_id = f"{uuid.uuid4().hex}{ext}"
    image_path = _job_image_path(job_id, image_id)
    image_file.save(str(image_path))

    try:
        with Image.open(str(image_path)) as img:
            width, height = img.size
    except Exception:
        image_path.unlink(missing_ok=True)
        return jsonify({"error": "Invalid image file"}), 400

    return jsonify({
        "ok": True,
        "image_id": image_id,
        "image_url": _job_image_url(job_id, image_id),
        "width": width,
        "height": height,
        "name": Path(image_file.filename).name,
    })


@app.route("/result/<job_id>/image/<image_id>")
def serve_translation_image(job_id: str, image_id: str):
    image_path = _job_image_path(job_id, image_id)
    if not image_path.exists():
        return jsonify({"error": "Image not found"}), 404
    return send_file(str(image_path))


@app.route("/pdf/<job_id>")
def serve_pdf(job_id: str):
    pdf_path = OUTPUT_FOLDER / f"{job_id}.pdf"
    if not pdf_path.exists():
        return jsonify({"error": "PDF not found"}), 404
    return send_file(str(pdf_path), mimetype="application/pdf")


@app.route("/word-meaning", methods=["POST"])
def word_meaning():
    body = request.get_json(force=True)
    word = (body.get("word") or "").strip()
    context = (body.get("context") or "").strip()
    target_lang = (body.get("target_lang") or TARGET_LANGUAGE).strip()
    api_key = (body.get("api_key") or GEMINI_API_KEY).strip()
    gemini_model = (body.get("gemini_model") or GEMINI_MODEL).strip()

    if not word:
        return jsonify({"error": "Missing word"}), 400
    if not api_key:
        return jsonify({"error": "Missing API Key"}), 400

    try:
        client = _get_client(api_key)
        prompt = (
            f"You are a bilingual dictionary assistant. "
            f"Explain the meaning of the word or phrase in {target_lang} concisely. "
            f"Include: 1) meaning in context, 2) part of speech, 3) a short example sentence. "
            f"Keep it brief (under 80 words). Respond in {target_lang}.\n\n"
            f"Word: {word}\nContext: {context[:500]}"
        )
        meaning = _call_gemini(client, "gemini-2.5-flash-lite", prompt)
        return jsonify({"meaning": meaning})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/download/<job_id>")
def download(job_id: str):
    mode = request.args.get("mode", "compare").strip().lower()
    if mode not in {"compare", "reading", "original_layout", "reading_compare"}:
        return jsonify({"error": "Unsupported download mode"}), 400
    try:
        font_size = max(8, min(24, int(request.args.get("font_size", "10"))))
    except (ValueError, TypeError):
        font_size = 10

    with STORE_LOCK:
        raw_data = deepcopy(job_results.get(job_id))
    if raw_data is None:
        return jsonify({"error": "Result not found"}), 404

    source_stem = Path(raw_data.get("source_name") or "translated_output").stem
    pdf_path = OUTPUT_FOLDER / f"{job_id}.pdf"
    if mode in {"original_layout", "reading_compare"} and not pdf_path.exists():
        return jsonify({"error": "Original PDF not found"}), 404

    suffix_by_mode = {
        "compare": "compare",
        "reading": "reading",
        "original_layout": "original_layout",
        "reading_compare": "reading_compare",
    }
    download_name_by_mode = {
        "compare": f"{source_stem}_compare.docx",
        "reading": f"{source_stem}_reading.docx",
        "original_layout": f"{source_stem}_original_layout.docx",
        "reading_compare": f"{source_stem}_reading_compare.docx",
    }
    output_path = OUTPUT_FOLDER / f"{job_id}_{suffix_by_mode[mode]}.docx"
    _build_word(
        raw_data.get("original", []),
        raw_data.get("translated", []),
        output_path,
        translated_blocks=raw_data.get("translated_blocks"),
        job_id=job_id,
        mode=mode,
        pdf_path=pdf_path,
        font_size=font_size,
    )
    return send_file(
        str(output_path),
        as_attachment=True,
        download_name=download_name_by_mode[mode],
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )



if __name__ == "__main__":
    url = "http://127.0.0.1:5000"
    if not _acquire_instance_lock():
        print(f"ButterLayer is already running at {url}")
        try:
            webbrowser.open(url)
        except Exception:
            pass
        sys.exit(0)

    UPLOAD_FOLDER.mkdir(parents=True, exist_ok=True)
    OUTPUT_FOLDER.mkdir(parents=True, exist_ok=True)
    _load_saved_results()
    _create_desktop_shortcut()
    print(f"ButterLayer running at {url}")
    if _FROZEN:
        threading.Timer(1.5, lambda: webbrowser.open(url)).start()
    app.run(debug=False, host="127.0.0.1", port=5000)
